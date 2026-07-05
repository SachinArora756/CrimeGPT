from datetime import datetime
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from app.documents.registry import TemplateDefinition, TemplateSection, SectionType


def render_docx(template: TemplateDefinition, data: dict, file_path: str):
    doc = DocxDocument()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    _render_header(doc, template)

    for sec in template.sections:
        if sec.section_type == SectionType.METADATA_TABLE:
            _render_metadata_table(doc, sec, data)
        elif sec.section_type == SectionType.BODY_TEXT:
            _render_body_text(doc, sec, data)
        elif sec.section_type == SectionType.LIST:
            _render_list(doc, sec, data)
        elif sec.section_type == SectionType.LEGAL_FOOTER:
            _render_legal_footer(doc, sec, data)

    _render_signatures(doc, template)
    doc.save(file_path)


def _render_header(doc: DocxDocument, template: TemplateDefinition):
    if template.form_number:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(template.form_number)
        run.font.size = Pt(9)
        run.italic = True

    heading = doc.add_heading(template.title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.size = Pt(16)

    if template.subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(template.subtitle)
        run.font.size = Pt(10)
        run.italic = True

    doc.add_paragraph("")


def _render_metadata_table(doc: DocxDocument, section: TemplateSection, data: dict):
    if section.title:
        h = doc.add_heading(section.title, level=2)
        for run in h.runs:
            run.font.size = Pt(12)

    table = doc.add_table(rows=len(section.fields), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, field_def in enumerate(section.fields):
        row = table.rows[i]
        label_cell = row.cells[0]
        value_cell = row.cells[1]

        label_cell.text = field_def.label
        value = data.get(field_def.key, field_def.fallback)
        value_cell.text = str(value) if value else field_def.fallback

        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

        label_cell.width = Inches(2.5)
        value_cell.width = Inches(4.0)

    doc.add_paragraph("")


def _render_body_text(doc: DocxDocument, section: TemplateSection, data: dict):
    if section.title:
        h = doc.add_heading(section.title, level=2)
        for run in h.runs:
            run.font.size = Pt(12)

    content = data.get(section.content_key, "---") if section.content_key else "---"
    if content and content != "---":
        for para_text in str(content).split("\n"):
            if para_text.strip():
                p = doc.add_paragraph(para_text.strip())
                p.paragraph_format.space_after = Pt(6)
    else:
        doc.add_paragraph("---")

    doc.add_paragraph("")


def _render_list(doc: DocxDocument, section: TemplateSection, data: dict):
    if section.title:
        h = doc.add_heading(section.title, level=2)
        for run in h.runs:
            run.font.size = Pt(12)

    items = data.get(section.content_key, []) if section.content_key else []
    if not items:
        doc.add_paragraph("(None)")
    else:
        for item in items:
            p = doc.add_paragraph(str(item))
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(3)

    doc.add_paragraph("")


def _render_legal_footer(doc: DocxDocument, section: TemplateSection, data: dict):
    content = data.get(section.content_key, "") if section.content_key else ""
    if content:
        doc.add_paragraph("")
        p = doc.add_paragraph()
        run = p.add_run(content)
        run.font.size = Pt(9)
        run.italic = True


def _render_signatures(doc: DocxDocument, template: TemplateDefinition):
    doc.add_paragraph("")
    doc.add_paragraph("")

    for sig in template.signatures:
        doc.add_paragraph("")
        p = doc.add_paragraph()
        p.add_run("_" * 40)
        p2 = doc.add_paragraph()
        run = p2.add_run(sig.title)
        run.font.size = Pt(10)

        if sig.with_date:
            p3 = doc.add_paragraph()
            run = p3.add_run(f"Date: {datetime.utcnow().strftime('%d/%m/%Y')}")
            run.font.size = Pt(9)

        if sig.with_seal:
            p4 = doc.add_paragraph()
            run = p4.add_run("[SEAL]")
            run.font.size = Pt(9)
            run.italic = True

    if template.seal_placeholder:
        doc.add_paragraph("")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run("[Official Seal / Stamp]")
        run.font.size = Pt(9)
        run.italic = True
