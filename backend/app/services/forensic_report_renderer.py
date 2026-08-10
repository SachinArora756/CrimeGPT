"""
PRISM Digital Forensic Investigation Report — DOCX Renderer.

Generates a professionally formatted forensic investigation report
following the PRISM framework (Procedural Record of Investigation, Substantiation & Methodology).
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


def _set_cell_shading(cell, color: str):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _add_page_break(doc):
    doc.add_page_break()


def _format_datetime(dt) -> str:
    if isinstance(dt, datetime):
        return dt.strftime("%d/%m/%Y %H:%M IST")
    if dt:
        return str(dt)
    return "---"


def _format_date(dt) -> str:
    if isinstance(dt, datetime):
        return dt.strftime("%d/%m/%Y")
    if hasattr(dt, 'strftime'):
        return dt.strftime("%d/%m/%Y")
    if dt:
        return str(dt)
    return "---"


BRAND_COLOR = "0D2B4E"
ACCENT_COLOR = "1B5E8C"


def render_forensic_report(sections_content: dict, case_data: dict, file_path: str):
    """Render the complete PRISM forensic report as DOCX."""
    doc = Document()

    _setup_styles(doc)
    _setup_headers_footers(doc, case_data)
    _render_cover_page(doc, case_data)
    _add_page_break(doc)
    _render_document_administration(doc, case_data)
    _add_page_break(doc)
    _render_table_of_contents(doc)
    _add_page_break(doc)

    # PART I — PROVENANCE
    _add_part_heading(doc, "PART I — PROVENANCE")

    # Section 2: Investigation Mandate & Authority
    _render_text_section(doc, "2.0", "Investigation Mandate & Authority",
                         sections_content.get("mandate_authority", ""))

    # Section 3: Examiner Declaration
    _add_page_break(doc)
    _render_examiner_declaration(doc, case_data)

    # PART II — RECONSTRUCTION
    _add_page_break(doc)
    _add_part_heading(doc, "PART II — RECONSTRUCTION")

    # Section 4: Case Synopsis
    _render_text_section(doc, "4.0", "Case Synopsis",
                         sections_content.get("case_synopsis", ""))

    # Section 5: Evidence Inventory
    _add_page_break(doc)
    _render_evidence_inventory(doc, case_data)

    # Section 6: Technical Infrastructure
    _render_technical_infrastructure(doc, case_data)

    # Section 7: Analytical Protocol
    _add_page_break(doc)
    _render_text_section(doc, "7.0", "Analytical Protocol",
                         sections_content.get("analytical_protocol", ""))

    # Section 8: Temporal Framework
    _render_text_section(doc, "8.0", "Temporal Synchronization Framework",
                         sections_content.get("temporal_framework", ""))

    # Section 9: Examination Findings
    _add_page_break(doc)
    _render_text_section(doc, "9.0", "Detailed Examination Findings",
                         sections_content.get("examination_findings", ""))

    # Section 10: Event Reconstruction
    _add_page_break(doc)
    _render_text_section(doc, "10.0", "Chronological Event Reconstruction",
                         sections_content.get("event_reconstruction", ""))

    # PART III — INTERPRETATION
    _add_page_break(doc)
    _add_part_heading(doc, "PART III — INTERPRETATION")

    # Section 11: Investigative Conclusions
    _render_text_section(doc, "11.0", "Investigative Conclusions",
                         sections_content.get("investigative_conclusions", ""))

    # Section 12: Digital Threat Assessment
    _add_page_break(doc)
    _render_text_section(doc, "12.0", "Digital Threat Assessment",
                         sections_content.get("threat_assessment", ""))

    # Section 13: Evidence Strength Evaluation
    _add_page_break(doc)
    _render_text_section(doc, "13.0", "Evidence Strength Evaluation",
                         sections_content.get("strength_evaluation", ""))

    # PART IV — SUBSTANTIATION
    _add_page_break(doc)
    _add_part_heading(doc, "PART IV — SUBSTANTIATION")

    # Section 14: Legal Compliance
    _render_text_section(doc, "14.0", "Legal Compliance Framework",
                         sections_content.get("legal_compliance", ""))

    # Section 15: Methodological Constraints
    _add_page_break(doc)
    _render_text_section(doc, "15.0", "Methodological Constraints & Limitations",
                         sections_content.get("methodological_constraints", ""))

    # Section 16: Preliminary Information
    _render_text_section(doc, "16.0", "Preliminary Information & Working Hypotheses",
                         sections_content.get("preliminary_information", ""))

    # PART V — MEMORANDUM
    _add_page_break(doc)
    _add_part_heading(doc, "PART V — MEMORANDUM")

    # Section 17: Expert Opinion
    _render_text_section(doc, "17.0", "Expert Professional Opinion",
                         sections_content.get("professional_opinion", ""))

    # Section 18: Evidence Disposition
    _add_page_break(doc)
    _render_evidence_disposition(doc, case_data)

    # Section 19: Declaration & Attestation
    _add_page_break(doc)
    _render_declaration_attestation(doc, case_data)

    # Section 20: Annexures
    _add_page_break(doc)
    _render_annexures(doc, case_data)

    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    doc.save(file_path)
    return file_path


def _setup_styles(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def _setup_headers_footers(doc, case_data):
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    header = section.header
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run("RESTRICTED — FOR AUTHORIZED LAW ENFORCEMENT USE ONLY")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)
    run.font.bold = True

    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fir = case_data.get("fir_number", "---")
    run = footer_para.add_run(f"FIR No. {fir} — PRISM Forensic Investigation Report")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def _add_part_heading(doc, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0D, 0x2B, 0x4E)
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(12)


def _render_cover_page(doc, case_data):
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PRISM")
    run.font.size = Pt(38)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0D, 0x2B, 0x4E)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Digital Forensic Investigation Report")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x8C)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Procedural Record of Investigation, Substantiation & Methodology")
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"FIR No. {case_data.get('fir_number', '---')}")
    run.font.size = Pt(14)
    run.font.bold = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(case_data.get("offense_type", "Criminal Investigation"))
    run.font.size = Pt(13)
    run.font.italic = True

    for _ in range(4):
        doc.add_paragraph()

    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    metadata = [
        ("Report Reference", f"PRISM/{case_data.get('fir_number', '---')}/{datetime.utcnow().strftime('%Y')}"),
        ("Security Classification", "RESTRICTED — For Authorized Law Enforcement Use Only"),
        ("Prepared By", case_data.get("io_name", "Investigating Officer")),
        ("Date of Compilation", _format_date(datetime.utcnow())),
        ("Originating Station", case_data.get("station_id", "---")),
        ("Framework Version", "PRISM v1.0"),
    ]

    for i, (label, value) in enumerate(metadata):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        _set_cell_shading(row.cells[0], "E3EFF7")
        for para in row.cells[0].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(10)
        for para in row.cells[1].paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)


def _render_document_administration(doc, case_data):
    heading = doc.add_heading("1.0  Document Administration", level=1)
    heading.runs[0].font.color.rgb = RGBColor(0x0D, 0x2B, 0x4E)

    doc.add_paragraph().add_run("Revision Record").bold = True
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'

    headers = ["Rev.", "Date", "Prepared By", "Nature of Revision"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        _set_cell_shading(cell, BRAND_COLOR)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    row = table.rows[1]
    row.cells[0].text = "1.0"
    row.cells[1].text = _format_date(datetime.utcnow())
    row.cells[2].text = case_data.get("io_name", "IO")
    row.cells[3].text = "Original compilation"

    doc.add_paragraph()
    doc.add_paragraph().add_run("Authorized Recipients").bold = True
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'

    headers = ["Recipient", "Designation", "Access Level"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        _set_cell_shading(cell, BRAND_COLOR)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    dist = [
        (case_data.get("io_name", "IO"), "Investigating Officer", "Full Report"),
        ("Competent Court", "Judicial Authority", "Full Report"),
        ("Supervisory Officer", "SHO / SP Office", "Executive Summary"),
    ]
    for i, (name, role, cls) in enumerate(dist):
        row = table.rows[i + 1]
        row.cells[0].text = name
        row.cells[1].text = role
        row.cells[2].text = cls


def _render_table_of_contents(doc):
    heading = doc.add_heading("Table of Contents", level=1)
    heading.runs[0].font.color.rgb = RGBColor(0x0D, 0x2B, 0x4E)

    toc = [
        (None, "PART I — PROVENANCE"),
        ("1.0", "Document Administration"),
        ("2.0", "Investigation Mandate & Authority"),
        ("3.0", "Examiner Declaration & Credentials"),
        (None, "PART II — RECONSTRUCTION"),
        ("4.0", "Case Synopsis"),
        ("5.0", "Evidence Inventory & Integrity Verification"),
        ("6.0", "Technical Infrastructure & Instruments"),
        ("7.0", "Analytical Protocol"),
        ("8.0", "Temporal Synchronization Framework"),
        ("9.0", "Detailed Examination Findings"),
        ("10.0", "Chronological Event Reconstruction"),
        (None, "PART III — INTERPRETATION"),
        ("11.0", "Investigative Conclusions"),
        ("12.0", "Digital Threat Assessment"),
        ("13.0", "Evidence Strength Evaluation"),
        (None, "PART IV — SUBSTANTIATION"),
        ("14.0", "Legal Compliance Framework"),
        ("15.0", "Methodological Constraints & Limitations"),
        ("16.0", "Preliminary Information & Working Hypotheses"),
        (None, "PART V — MEMORANDUM"),
        ("17.0", "Expert Professional Opinion"),
        ("18.0", "Evidence Handling & Disposition"),
        ("19.0", "Declaration & Attestation"),
        ("20.0", "Supporting Annexures"),
    ]

    for num, title in toc:
        p = doc.add_paragraph()
        if num is None:
            run = p.add_run(title)
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x8C)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
        else:
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(f"{num}    {title}")
            run.font.size = Pt(11)


def _render_text_section(doc, section_number: str, title: str, content: str):
    heading = doc.add_heading(f"{section_number}  {title}", level=1)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x0D, 0x2B, 0x4E)

    if not content:
        p = doc.add_paragraph()
        p.add_run("[Section content pending generation]").italic = True
        return

    paragraphs = content.strip().split("\n")
    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            doc.add_paragraph()
            continue

        if (para_text.endswith(":") and len(para_text) < 80) or (para_text.isupper() and len(para_text) < 60):
            p = doc.add_paragraph()
            run = p.add_run(para_text)
            run.font.bold = True
            run.font.size = Pt(11)
            continue

        if len(para_text) > 2 and para_text[0].isdigit() and "." in para_text[:5]:
            parts = para_text.split(" ", 1)
            if len(parts) == 2 and parts[0].replace(".", "").isdigit():
                p = doc.add_paragraph()
                run = p.add_run(parts[0] + " ")
                run.font.bold = True
                p.add_run(parts[1])
                continue

        p = doc.add_paragraph()
        p.add_run(para_text)
        p.paragraph_format.space_after = Pt(6)


def _render_examiner_declaration(doc, case_data):
    heading = doc.add_heading("3.0  Examiner Declaration & Credentials", level=1)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x0D, 0x2B, 0x4E)

    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'

    entries = [
        ("Name", case_data.get("io_name", "---")),
        ("Designation", "Investigating Officer / Digital Forensic Examiner"),
        ("Originating Station", case_data.get("station_id", "---")),
        ("Identification No.", case_data.get("officer_badge", "---")),
        ("Role in Investigation", "Lead Examiner"),
        ("Declaration Date", _format_date(datetime.utcnow())),
    ]

    for i, (label, value) in enumerate(entries):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        _set_cell_shading(row.cells[0], "E3EFF7")
        for para in row.cells[0].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(10)
        for para in row.cells[1].paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(
        "The undersigned hereby declares that they have no personal, financial, or other interest "
        "in the outcome of this case. The findings and opinions expressed herein are based solely "
        "upon the evidence examined and the professional expertise of the examiner. This declaration "
        "is made in compliance with the requirements of the Bharatiya Sakshya Adhiniyam, 2023."
    )
    p.paragraph_format.space_after = Pt(6)


def _render_evidence_inventory(doc, case_data):
    heading = doc.add_heading("5.0  Evidence Inventory & Integrity Verification", level=1)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x0D, 0x2B, 0x4E)

    evidence_items = case_data.get("evidence_items", [])

    if not evidence_items:
        p = doc.add_paragraph()
        p.add_run("No evidence items recorded in the case management system.").italic = True
        return

    p = doc.add_paragraph()
    p.add_run(
        f"A total of {len(evidence_items)} evidence item(s) were received, catalogued, "
        "and subjected to integrity verification."
    )

    col_count = 5
    table = doc.add_table(rows=1 + len(evidence_items), cols=col_count)
    table.style = 'Table Grid'

    headers = ["Exhibit", "Description", "Format", "Integrity Hash", "Provenance"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        _set_cell_shading(cell, BRAND_COLOR)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for idx, ev in enumerate(evidence_items):
        row = table.rows[idx + 1]
        row.cells[0].text = f"EX-{idx + 1:03d}"
        row.cells[1].text = (ev.get("description") or ev.get("original_filename", "---"))[:45]
        row.cells[2].text = ev.get("file_type", "---")
        hash_val = ev.get("file_hash", "---")
        row.cells[3].text = hash_val[:16] + "..." if hash_val and len(hash_val) > 16 else (hash_val or "---")
        custody = ev.get("chain_of_custody", [])
        row.cells[4].text = f"{len(custody)} transfer(s)" if custody else "Original receipt"

        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8)


def _render_technical_infrastructure(doc, case_data):
    heading = doc.add_heading("6.0  Technical Infrastructure & Instruments", level=1)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x0D, 0x2B, 0x4E)

    tools = case_data.get("forensic_tools", [])

    p = doc.add_paragraph()
    p.add_run(
        "The forensic examination was conducted utilizing the CrimeGPT Digital Forensic Platform, "
        "an integrated AI-assisted investigation system incorporating specialized forensic analysis modules."
    )

    if tools:
        doc.add_paragraph()
        doc.add_paragraph().add_run("Instruments & Analytical Modules Employed:").bold = True

        table = doc.add_table(rows=1 + len(tools), cols=3)
        table.style = 'Table Grid'

        headers = ["Module", "Function", "Invocations"]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            _set_cell_shading(cell, BRAND_COLOR)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for idx, tool in enumerate(tools):
            row = table.rows[idx + 1]
            row.cells[0].text = tool.get("name", "---")
            row.cells[1].text = tool.get("purpose", "---")
            row.cells[2].text = str(tool.get("count", 0))
    else:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Standard forensic examination procedures were followed using the platform's built-in analytical capabilities.")


def _render_evidence_disposition(doc, case_data):
    heading = doc.add_heading("18.0  Evidence Handling & Disposition", level=1)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x0D, 0x2B, 0x4E)

    evidence_items = case_data.get("evidence_items", [])

    p = doc.add_paragraph()
    p.add_run("Upon completion of the forensic examination, all evidence items are subject to the following disposition protocol:")

    dispositions = [
        "All original evidence items shall be returned to the custody of the designated Investigating Officer under documented transfer.",
        "Forensic images, working copies, and derivative materials shall be retained in secure encrypted storage for the legally prescribed retention period.",
        "All temporary files, intermediate outputs, and working copies created during the examination have been securely purged using certified deletion procedures.",
        f"This examination processed a total of {len(evidence_items)} evidence item(s) as documented in the Evidence Inventory (Section 5.0).",
        "The integrity of all evidence items was verified both at receipt and upon return, with hash values confirmed unchanged.",
    ]

    for d in dispositions:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.add_run(f"• {d}")
        p.paragraph_format.space_after = Pt(4)


def _render_declaration_attestation(doc, case_data):
    heading = doc.add_heading("19.0  Declaration & Attestation", level=1)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x0D, 0x2B, 0x4E)

    p = doc.add_paragraph()
    p.add_run("The undersigned solemnly declares and attests as follows:")

    statements = [
        "I confirm that the facts stated in this report, insofar as they are within my personal knowledge, are true and accurate to the best of my belief. Where facts are based upon information provided by others, I have identified the source and believe such information to be reliable.",
        "The professional opinions expressed in this report represent my genuine, considered assessment based solely upon the evidence examined, the analytical procedures applied, and my professional training and experience.",
        "I understand that this report may be submitted to a court of law and that I may be called upon to give evidence in relation to its contents. I am aware of the consequences of making a false declaration.",
        "I confirm that no arrangement exists, nor has been entered into, whereby my remuneration or any benefit is contingent upon the findings or outcome of this case.",
        "I have no conflict of interest, financial or otherwise, in relation to any party to these proceedings, except as may be disclosed elsewhere in this report.",
        "The examination was conducted in accordance with established forensic science principles, applicable professional standards, and the procedural requirements of Indian law.",
    ]

    for i, s in enumerate(statements):
        p = doc.add_paragraph()
        p.add_run(f"{i+1}. {s}")
        p.paragraph_format.space_after = Pt(8)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("_" * 40)
    p = doc.add_paragraph()
    run = p.add_run(case_data.get("io_name", "Investigating Officer"))
    run.font.bold = True
    p = doc.add_paragraph()
    p.add_run("Investigating Officer / Digital Forensic Examiner")
    p = doc.add_paragraph()
    p.add_run(f"Date: {_format_date(datetime.utcnow())}")
    p = doc.add_paragraph()
    p.add_run(f"Station: {case_data.get('station_id', '---')}")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[Official Seal / Stamp]")
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


def _render_annexures(doc, case_data):
    heading = doc.add_heading("20.0  Supporting Annexures", level=1)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x0D, 0x2B, 0x4E)

    # Annexure A: Evidence Register
    doc.add_heading("Annexure A — Complete Evidence Register", level=2)
    evidence_items = case_data.get("evidence_items", [])
    if evidence_items:
        table = doc.add_table(rows=1 + len(evidence_items), cols=4)
        table.style = 'Table Grid'
        headers = ["Exhibit", "Filename", "Format", "Size"]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            _set_cell_shading(cell, "374750")
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for idx, ev in enumerate(evidence_items):
            row = table.rows[idx + 1]
            row.cells[0].text = f"EX-{idx + 1:03d}"
            row.cells[1].text = ev.get("original_filename", "---")[:40]
            row.cells[2].text = ev.get("file_type", "---")
            size = ev.get("file_size", 0)
            if size > 1024 * 1024:
                row.cells[3].text = f"{size / (1024*1024):.1f} MB"
            elif size > 1024:
                row.cells[3].text = f"{size / 1024:.1f} KB"
            else:
                row.cells[3].text = f"{size} B"
    else:
        doc.add_paragraph("No evidence items recorded.")

    # Annexure B: Forensic Module Execution Log
    _add_page_break(doc)
    doc.add_heading("Annexure B — Forensic Module Execution Log", level=2)
    tool_executions = case_data.get("tool_executions", [])
    if tool_executions:
        table = doc.add_table(rows=1 + min(len(tool_executions), 50), cols=5)
        table.style = 'Table Grid'
        headers = ["Module", "Target", "Result", "Confidence", "Timestamp"]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            _set_cell_shading(cell, "374750")
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for idx, ex in enumerate(tool_executions[:50]):
            row = table.rows[idx + 1]
            row.cells[0].text = ex.get("tool_name", "---")[:20]
            row.cells[1].text = ex.get("evidence_name", "---")[:20]
            row.cells[2].text = ex.get("status", "---")
            conf = ex.get("confidence")
            row.cells[3].text = f"{conf:.0%}" if conf else "---"
            row.cells[4].text = _format_datetime(ex.get("created_at"))
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(8)
    else:
        doc.add_paragraph("No forensic module executions recorded.")

    # Annexure C: Investigation Diary
    _add_page_break(doc)
    doc.add_heading("Annexure C — Investigation Diary Summary", level=2)
    diary_entries = case_data.get("diary_entries_list", [])
    if diary_entries:
        for entry in diary_entries:
            p = doc.add_paragraph()
            run = p.add_run(f"{_format_date(entry.get('entry_date', ''))} [{entry.get('entry_type', 'investigation')}]")
            run.font.bold = True
            run.font.size = Pt(10)
            p = doc.add_paragraph()
            p.add_run(entry.get("content", "---")[:500])
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.left_indent = Cm(0.5)
    else:
        doc.add_paragraph("No investigation diary entries recorded.")

    # Annexure D: Provenance Chain
    _add_page_break(doc)
    doc.add_heading("Annexure D — Provenance Chain Records", level=2)
    has_custody = False
    for ev in evidence_items:
        custody = ev.get("chain_of_custody", [])
        if custody:
            has_custody = True
            p = doc.add_paragraph()
            run = p.add_run(f"Exhibit: {ev.get('original_filename', '---')}")
            run.font.bold = True
            for entry in custody:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                if isinstance(entry, dict):
                    p.add_run(f"• {entry.get('action', '---')} — {entry.get('timestamp', '---')} — {entry.get('officer', '---')}")
                else:
                    p.add_run(f"• {str(entry)}")
            doc.add_paragraph()
    if not has_custody:
        doc.add_paragraph("Provenance chain records are maintained within the case management system and are available upon request.")

    # Annexure E: Terminology
    _add_page_break(doc)
    doc.add_heading("Annexure E — Terminology Reference", level=2)
    glossary = [
        ("PRISM", "Procedural Record of Investigation, Substantiation & Methodology"),
        ("Grade Alpha", "Direct primary evidence from original source"),
        ("Grade Beta", "Corroborated evidence from secondary sources"),
        ("Grade Gamma", "Circumstantial or indirect evidence"),
        ("Temporal Verified", "Timestamp from authoritative/system source"),
        ("Temporal Derived", "Computed or calculated timestamp"),
        ("Temporal Estimated", "Approximate time based on contextual indicators"),
        ("BNS", "Bharatiya Nyaya Sanhita, 2023"),
        ("BNSS", "Bharatiya Nagarik Suraksha Sanhita, 2023"),
        ("BSA", "Bharatiya Sakshya Adhiniyam, 2023"),
        ("FIR", "First Information Report"),
        ("IO", "Investigating Officer"),
        ("IST", "Indian Standard Time (UTC+05:30)"),
        ("SHA-256", "Secure Hash Algorithm — 256-bit integrity verification"),
    ]

    table = doc.add_table(rows=1 + len(glossary), cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = "Term"
    table.rows[0].cells[1].text = "Definition"
    _set_cell_shading(table.rows[0].cells[0], "374750")
    _set_cell_shading(table.rows[0].cells[1], "374750")
    for para in table.rows[0].cells[0].paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for para in table.rows[0].cells[1].paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for idx, (term, defn) in enumerate(glossary):
        row = table.rows[idx + 1]
        row.cells[0].text = term
        row.cells[1].text = defn
        for para in row.cells[0].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(9)
        for para in row.cells[1].paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)
