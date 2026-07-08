"""
Digital Forensics Toolkit - Tool Execution Handlers

Each handler accepts (file_path: str, params: dict) and returns
(output_data: dict, confidence: float | None).

Blocking operations are wrapped in asyncio.to_thread().
Import errors are handled gracefully with descriptive messages.
"""

import os
import hashlib
import json
import asyncio
import re
import time
import logging
from datetime import datetime
from pathlib import Path
from typing import Any

import numpy as np

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Singleton model loaders (lazy)
# ---------------------------------------------------------------------------

_insightface_app = None
_clip_model = None
_clip_preprocess = None
_clip_tokenizer = None


def _get_insightface():
    global _insightface_app
    if _insightface_app is None:
        from insightface.app import FaceAnalysis
        _insightface_app = FaceAnalysis(
            name="buffalo_l",
            providers=["CPUExecutionProvider"],
        )
        _insightface_app.prepare(ctx_id=0, det_size=(640, 640))
    return _insightface_app


def _get_clip_model():
    global _clip_model, _clip_preprocess
    if _clip_model is None:
        import open_clip
        _clip_model, _, _clip_preprocess = open_clip.create_model_and_transforms(
            "ViT-B-32", pretrained="laion2b_s34b_b79k"
        )
        _clip_model.eval()
    return _clip_model, _clip_preprocess


# ---------------------------------------------------------------------------
# Qdrant helpers
# ---------------------------------------------------------------------------

FACE_COLLECTION = "face_embeddings"
IMAGE_COLLECTION = "image_embeddings"
FACE_EMBEDDING_DIM = 512
CLIP_EMBEDDING_DIM = 512


def _get_qdrant_client():
    from qdrant_client import QdrantClient
    from app.config import settings
    return QdrantClient(host=settings.qdrant_host, port=settings.qdrant_port)


def _ensure_qdrant_collection(client, collection_name: str, dim: int):
    from qdrant_client.models import VectorParams, Distance
    collections = [c.name for c in client.get_collections().collections]
    if collection_name not in collections:
        client.create_collection(
            collection_name=collection_name,
            vectors_config=VectorParams(size=dim, distance=Distance.COSINE),
        )


# ---------------------------------------------------------------------------
# Database helper (sync context for to_thread)
# ---------------------------------------------------------------------------

def _get_sync_db_url():
    from app.config import settings
    url = settings.database_url
    return url.replace("postgresql+asyncpg://", "postgresql://")


def _query_criminal_profiles_sync(criminal_ids: list[int]) -> list[dict]:
    """Query criminal profiles by internal IDs using a sync connection."""
    import sqlalchemy
    from sqlalchemy import create_engine, text

    engine = create_engine(_get_sync_db_url(), pool_pre_ping=True)
    results = []
    with engine.connect() as conn:
        if not criminal_ids:
            return results
        placeholders = ",".join(str(int(cid)) for cid in criminal_ids)
        query = text(f"""
            SELECT id, criminal_id, full_name, father_name, gender,
                   wanted_status, danger_level, gang_name, total_arrests,
                   total_firs, crime_categories, identifying_marks
            FROM criminal_profiles
            WHERE id IN ({placeholders}) AND is_active = true
        """)
        rows = conn.execute(query).fetchall()
        for row in rows:
            results.append({
                "id": row[0],
                "criminal_id": row[1],
                "full_name": row[2],
                "father_name": row[3],
                "gender": row[4],
                "wanted_status": row[5],
                "danger_level": row[6],
                "gang_name": row[7],
                "total_arrests": row[8],
                "total_firs": row[9],
                "crime_categories": row[10],
                "identifying_marks": row[11],
            })
    engine.dispose()
    return results


def _query_all_face_embeddings_sync() -> list[dict]:
    """Fetch all face embeddings from DB for comparison."""
    import sqlalchemy
    from sqlalchemy import create_engine, text

    engine = create_engine(_get_sync_db_url(), pool_pre_ping=True)
    results = []
    with engine.connect() as conn:
        query = text("""
            SELECT fe.id, fe.criminal_id, fe.embedding, fe.image_path, fe.quality_score,
                   cp.full_name, cp.criminal_id as crim_public_id, cp.wanted_status,
                   cp.danger_level, cp.gang_name
            FROM criminal_face_embeddings fe
            JOIN criminal_profiles cp ON fe.criminal_id = cp.id
            WHERE cp.is_active = true
        """)
        rows = conn.execute(query).fetchall()
        for row in rows:
            results.append({
                "db_id": row[0],
                "criminal_db_id": row[1],
                "embedding": row[2],
                "image_path": row[3],
                "quality_score": row[4],
                "full_name": row[5],
                "criminal_id": row[6],
                "wanted_status": row[7],
                "danger_level": row[8],
                "gang_name": row[9],
            })
    engine.dispose()
    return results


def _query_all_fingerprints_sync() -> list[dict]:
    """Fetch all fingerprint templates from DB."""
    import sqlalchemy
    from sqlalchemy import create_engine, text

    engine = create_engine(_get_sync_db_url(), pool_pre_ping=True)
    results = []
    with engine.connect() as conn:
        query = text("""
            SELECT fp.id, fp.criminal_id, fp.finger_type, fp.template_data,
                   fp.quality_score, cp.full_name, cp.criminal_id as crim_public_id,
                   cp.wanted_status, cp.danger_level
            FROM criminal_fingerprints fp
            JOIN criminal_profiles cp ON fp.criminal_id = cp.id
            WHERE cp.is_active = true
        """)
        rows = conn.execute(query).fetchall()
        for row in rows:
            results.append({
                "db_id": row[0],
                "criminal_db_id": row[1],
                "finger_type": row[2],
                "template_data": row[3],
                "quality_score": row[4],
                "full_name": row[5],
                "criminal_id": row[6],
                "wanted_status": row[7],
                "danger_level": row[8],
            })
    engine.dispose()
    return results


def _query_all_dna_profiles_sync() -> list[dict]:
    """Fetch all DNA profiles from DB."""
    import sqlalchemy
    from sqlalchemy import create_engine, text

    engine = create_engine(_get_sync_db_url(), pool_pre_ping=True)
    results = []
    with engine.connect() as conn:
        query = text("""
            SELECT dp.id, dp.criminal_id, dp.dna_id, dp.sample_number,
                   dp.laboratory, dp.loci_markers, cp.full_name,
                   cp.criminal_id as crim_public_id, cp.wanted_status, cp.danger_level
            FROM criminal_dna_profiles dp
            JOIN criminal_profiles cp ON dp.criminal_id = cp.id
            WHERE cp.is_active = true
        """)
        rows = conn.execute(query).fetchall()
        for row in rows:
            results.append({
                "db_id": row[0],
                "criminal_db_id": row[1],
                "dna_id": row[2],
                "sample_number": row[3],
                "laboratory": row[4],
                "loci_markers": row[5],
                "full_name": row[6],
                "criminal_id": row[7],
                "wanted_status": row[8],
                "danger_level": row[9],
            })
    engine.dispose()
    return results


async def _enrich_matches_with_profiles(matches: list[dict]) -> list[dict]:
    """Fetch full criminal profiles for matched criminal_ids and merge into match dicts."""
    if not matches:
        return matches

    import asyncpg

    criminal_ids = list({m.get("criminal_id", "") for m in matches if m.get("criminal_id")})
    if not criminal_ids:
        return matches

    from app.config import settings
    db_url = settings.database_url.replace("postgresql+asyncpg://", "postgresql://")
    profiles_map: dict[str, dict] = {}

    try:
        conn = await asyncpg.connect(db_url)
        try:
            rows = await conn.fetch("""
                SELECT criminal_id, full_name, father_name, gender, nationality,
                       date_of_birth, height_cm, weight_kg, build, complexion,
                       hair_color, eye_color, identifying_marks, crime_categories,
                       modus_operandi, gang_name, gang_role, known_weapons,
                       total_arrests, total_convictions, total_firs,
                       first_offense_date, last_known_activity, occupation,
                       education, reward_amount, bail_status, wanted_status,
                       danger_level, nicknames
                FROM criminal_profiles
                WHERE criminal_id = ANY($1) AND is_active = true
            """, criminal_ids)
            for row in rows:
                profiles_map[row["criminal_id"]] = {
                    "profile_full_name": row["full_name"],
                    "father_name": row["father_name"],
                    "gender": row["gender"],
                    "nationality": row["nationality"],
                    "date_of_birth": str(row["date_of_birth"]) if row["date_of_birth"] else None,
                    "height_cm": float(row["height_cm"]) if row["height_cm"] else None,
                    "weight_kg": float(row["weight_kg"]) if row["weight_kg"] else None,
                    "build": row["build"],
                    "complexion": row["complexion"],
                    "hair_color": row["hair_color"],
                    "eye_color": row["eye_color"],
                    "identifying_marks": row["identifying_marks"],
                    "crime_categories": row["crime_categories"],
                    "modus_operandi": row["modus_operandi"],
                    "gang_name": row["gang_name"],
                    "gang_role": row["gang_role"],
                    "known_weapons": row["known_weapons"],
                    "total_arrests": row["total_arrests"],
                    "total_convictions": row["total_convictions"],
                    "total_firs": row["total_firs"],
                    "first_offense_date": str(row["first_offense_date"]) if row["first_offense_date"] else None,
                    "last_known_activity": str(row["last_known_activity"]) if row["last_known_activity"] else None,
                    "occupation": row["occupation"],
                    "education": row["education"],
                    "reward_amount": float(row["reward_amount"]) if row["reward_amount"] else None,
                    "bail_status": row["bail_status"],
                    "wanted_status": row["wanted_status"],
                    "danger_level": row["danger_level"],
                    "nicknames": row["nicknames"],
                }
        finally:
            await conn.close()
    except Exception as e:
        logger.warning(f"Failed to enrich matches with profiles: {e}")

    for match in matches:
        cid = match.get("criminal_id", "")
        if cid in profiles_map:
            match["criminal_profile"] = profiles_map[cid]

    return matches


# ---------------------------------------------------------------------------
# Image Tools
# ---------------------------------------------------------------------------

async def run_image_ocr(file_path: str, params: dict) -> tuple[dict, float | None]:
    """Run OCR on image file using EasyOCR (primary) with Tesseract fallback."""

    def _ocr():
        lang = params.get("language", "eng")

        # Strategy 1: EasyOCR (deep learning — works on scene text in photos)
        try:
            import easyocr
            lang_map = {"eng": "en", "hin": "hi", "en": "en", "hi": "hi"}
            ocr_lang = lang_map.get(lang, "en")
            reader = easyocr.Reader([ocr_lang], gpu=False, verbose=False)
            results = reader.readtext(file_path)

            if results:
                text_parts = []
                total_conf = 0.0
                regions = []
                for (bbox, text, conf) in results:
                    text_parts.append(text)
                    total_conf += conf
                    (tl, tr, br, bl) = bbox
                    regions.append({
                        "text": text,
                        "confidence": round(conf, 3),
                        "bbox": {"x": int(tl[0]), "y": int(tl[1]),
                                 "width": int(tr[0] - tl[0]), "height": int(bl[1] - tl[1])},
                    })

                full_text = " ".join(text_parts)
                avg_conf = round(total_conf / len(results), 3)

                return {
                    "text": full_text.strip(),
                    "char_count": len(full_text.strip()),
                    "line_count": len(text_parts),
                    "language": lang,
                    "method": "easyocr",
                    "regions_detected": len(regions),
                    "regions": regions[:50],
                }, avg_conf
        except ImportError:
            pass
        except Exception:
            pass

        # Strategy 2: Tesseract (good for document scans)
        try:
            import pytesseract
            from PIL import Image, ImageEnhance
        except ImportError as e:
            return {"error": f"Required library not available: {e}", "text": ""}, None

        try:
            img = Image.open(file_path)
            enhanced = ImageEnhance.Contrast(img).enhance(1.3)
            text = pytesseract.image_to_string(enhanced, lang=lang)
            img.close()

            confidence = None
            try:
                data = pytesseract.image_to_data(Image.open(file_path), output_type=pytesseract.Output.DICT)
                confs = [int(c) for c in data["conf"] if str(c).isdigit() and int(c) > 0]
                if confs:
                    confidence = round(sum(confs) / len(confs) / 100.0, 3)
            except Exception:
                pass

            return {
                "text": text.strip(),
                "char_count": len(text.strip()),
                "line_count": len(text.strip().splitlines()),
                "language": lang,
                "method": "tesseract",
            }, confidence
        except Exception as e:
            return {"error": f"OCR failed: {str(e)}", "text": ""}, None

    return await asyncio.to_thread(_ocr)


async def run_object_detection(file_path: str, params: dict) -> tuple[dict, float | None]:
    """Run YOLO object detection on image with Gemini Vision enhancement."""

    def _detect_yolo():
        try:
            from ultralytics import YOLO
        except ImportError:
            return None

        try:
            model_path = params.get("model", "yolov8n.pt")
            confidence_threshold = float(params.get("confidence", 0.15))
            model = YOLO(model_path)
            results = model(file_path, conf=confidence_threshold)

            objects = []
            total_conf = 0.0

            for result in results:
                for box in result.boxes:
                    cls_id = int(box.cls[0])
                    cls_name = result.names[cls_id]
                    conf = float(box.conf[0])
                    x1, y1, x2, y2 = box.xyxy[0].tolist()
                    objects.append({
                        "class": cls_name,
                        "confidence": round(conf, 4),
                        "bbox": {
                            "x1": round(x1, 1),
                            "y1": round(y1, 1),
                            "x2": round(x2, 1),
                            "y2": round(y2, 1),
                        },
                    })
                    total_conf += conf

            avg_conf = round(total_conf / len(objects), 3) if objects else None
            return {
                "objects_detected": len(objects),
                "objects": objects,
                "model_used": model_path,
                "method": "yolo",
            }, avg_conf
        except Exception:
            return None

    def _detect_gemini():
        from app.ai.llm_provider import has_any_llm_key, generate_vision
        if not has_any_llm_key():
            return None
        try:
            prompt = (
                "List every distinct object visible in this image. For each object provide:\n"
                "OBJECT: <name> | <description>\n"
                "Be thorough — include vehicles, people, animals, furniture, signs, text, "
                "tools, electronics, clothing items, etc. One line per object."
            )
            response_text = generate_vision(file_path, prompt, temperature=0.1, max_tokens=500)

            objects = []
            for line in response_text.strip().split("\n"):
                line = line.strip()
                if line.startswith("OBJECT:"):
                    parts = [p.strip() for p in line[7:].split("|")]
                    name = parts[0] if parts else "unknown"
                    desc = parts[1] if len(parts) > 1 else ""
                    objects.append({
                        "class": name.lower(),
                        "description": desc,
                        "confidence": 0.85,
                        "method": "gemini_vision",
                    })

            return {
                "objects_detected": len(objects),
                "objects": objects,
                "method": "gemini_vision",
            }, 0.85 if objects else None
        except Exception:
            return None

    def _run():
        yolo_result = _detect_yolo()
        if yolo_result and yolo_result[0].get("objects_detected", 0) >= 2:
            return yolo_result

        gemini_result = _detect_gemini()
        if gemini_result is not None:
            if yolo_result and yolo_result[0].get("objects_detected", 0) > 0:
                combined = yolo_result[0].copy()
                combined["gemini_objects"] = gemini_result[0].get("objects", [])
                combined["method"] = "yolo + gemini_vision"
                return combined, yolo_result[1]
            return gemini_result

        if yolo_result is not None:
            return yolo_result
        return {"objects_detected": 0, "objects": [], "method": "none"}, None

    return await asyncio.to_thread(_run)


async def run_image_exif(file_path: str, params: dict) -> tuple[dict, float | None]:
    """Extract EXIF data from image."""

    def _exif():
        try:
            from PIL import Image
            from PIL.ExifTags import TAGS, GPSTAGS
        except ImportError:
            return {"error": "Pillow not installed"}, None

        try:
            img = Image.open(file_path)
            exif_raw = img._getexif()
            img_info = {
                "width": img.width,
                "height": img.height,
                "format": img.format,
                "mode": img.mode,
            }
            img.close()

            if not exif_raw:
                return {
                    "image_info": img_info,
                    "exif_available": False,
                    "exif_data": {},
                    "note": "No EXIF data found in this image",
                }, None

            exif_data = {}
            gps_data = {}
            for tag_id, value in exif_raw.items():
                tag_name = TAGS.get(tag_id, str(tag_id))
                if tag_name == "GPSInfo":
                    for gps_tag_id, gps_value in value.items():
                        gps_tag_name = GPSTAGS.get(gps_tag_id, str(gps_tag_id))
                        gps_data[gps_tag_name] = str(gps_value)
                else:
                    try:
                        if isinstance(value, bytes):
                            exif_data[tag_name] = value.hex()[:100]
                        else:
                            exif_data[tag_name] = str(value)[:500]
                    except Exception:
                        exif_data[tag_name] = "<unreadable>"

            result = {
                "image_info": img_info,
                "exif_available": True,
                "exif_tags_count": len(exif_data),
                "exif_data": exif_data,
            }
            if gps_data:
                result["gps_data"] = gps_data

            return result, 1.0
        except Exception as e:
            return {"error": f"EXIF extraction failed: {str(e)}"}, None

    return await asyncio.to_thread(_exif)


async def run_audio_transcribe(file_path: str, params: dict) -> tuple[dict, float | None]:
    """Transcribe audio file using Whisper (torch-based)."""

    def _transcribe():
        # Try faster-whisper first (if ctranslate2 works)
        try:
            from faster_whisper import WhisperModel
            model_size = params.get("model", "base")
            language = params.get("language", None)

            model = WhisperModel(model_size, device="cpu", compute_type="int8")

            transcribe_kwargs = {}
            if language:
                transcribe_kwargs["language"] = language

            segments_iter, info = model.transcribe(file_path, **transcribe_kwargs)

            segments = []
            full_text_parts = []
            for segment in segments_iter:
                segments.append({
                    "start": round(segment.start, 2),
                    "end": round(segment.end, 2),
                    "text": segment.text.strip(),
                })
                full_text_parts.append(segment.text.strip())
                if len(segments) >= 200:
                    break

            full_text = " ".join(full_text_parts)
            duration = segments[-1]["end"] if segments else 0

            return {
                "text": full_text,
                "language_detected": info.language,
                "language_probability": round(info.language_probability, 3),
                "duration_seconds": round(duration, 2),
                "segment_count": len(segments),
                "segments": segments[:50],
                "model_used": f"faster-whisper ({model_size})",
            }, round(info.language_probability, 3)
        except (ImportError, OSError):
            pass

        # Fallback: use transformers pipeline (torch-based Whisper)
        try:
            import torch
            from transformers import pipeline

            model_size = params.get("model", "base")
            model_id = f"openai/whisper-{model_size}"

            pipe = pipeline(
                "automatic-speech-recognition",
                model=model_id,
                device="cpu",
                torch_dtype=torch.float32,
            )

            transcribe_kwargs = {"return_timestamps": True}
            language = params.get("language", None)
            if language:
                transcribe_kwargs["generate_kwargs"] = {"language": language}

            result = pipe(file_path, **transcribe_kwargs)

            text = result.get("text", "").strip()
            chunks = result.get("chunks", [])

            segments = []
            for chunk in chunks[:50]:
                ts = chunk.get("timestamp", (0, 0))
                segments.append({
                    "start": round(ts[0], 2) if ts[0] else 0,
                    "end": round(ts[1], 2) if ts[1] else 0,
                    "text": chunk.get("text", "").strip(),
                })

            duration = segments[-1]["end"] if segments else 0

            return {
                "text": text,
                "language_detected": language or "auto",
                "duration_seconds": round(duration, 2),
                "segment_count": len(segments),
                "segments": segments,
                "model_used": f"transformers/whisper-{model_size}",
            }, 0.85
        except Exception as e:
            return {"error": f"Audio transcription failed: {str(e)}", "text": ""}, None

    return await asyncio.to_thread(_transcribe)


async def run_document_ocr(file_path: str, params: dict) -> tuple[dict, float | None]:
    """Run OCR on document/PDF pages using EasyOCR (primary) with Tesseract fallback."""

    def _doc_ocr():
        ext = os.path.splitext(file_path)[1].lower()
        lang = params.get("language", "eng")

        if ext == ".pdf":
            try:
                import fitz
                import pytesseract
                from PIL import Image
            except ImportError as e:
                return {"error": f"PDF OCR requires PyMuPDF and pytesseract: {e}", "text": ""}, None

            try:
                doc = fitz.open(file_path)
                full_text = []
                max_pages = int(params.get("max_pages", 20))
                total_pages = doc.page_count

                for page_num in range(min(total_pages, max_pages)):
                    page = doc[page_num]
                    page_text = page.get_text().strip()
                    if page_text:
                        full_text.append(f"--- Page {page_num + 1} ---\n{page_text}")
                    else:
                        pix = page.get_pixmap(dpi=200)
                        img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
                        text = pytesseract.image_to_string(img, lang=lang)
                        full_text.append(f"--- Page {page_num + 1} ---\n{text.strip()}")

                doc.close()
                combined = "\n\n".join(full_text)
                return {
                    "text": combined,
                    "pages_processed": min(total_pages, max_pages),
                    "total_pages": total_pages,
                    "char_count": len(combined),
                    "method": "pdf_text_extraction + tesseract_fallback",
                }, 0.8
            except Exception as e:
                return {"error": f"Document OCR failed: {str(e)}", "text": ""}, None

        # For image files — use EasyOCR first
        try:
            import easyocr
            lang_map = {"eng": "en", "hin": "hi", "en": "en", "hi": "hi"}
            ocr_lang = lang_map.get(lang, "en")
            reader = easyocr.Reader([ocr_lang], gpu=False, verbose=False)
            results = reader.readtext(file_path)
            if results:
                text_parts = [text for (_, text, _) in results]
                full_text = "\n".join(text_parts)
                avg_conf = sum(conf for (_, _, conf) in results) / len(results)
                return {
                    "text": full_text.strip(),
                    "char_count": len(full_text.strip()),
                    "line_count": len(text_parts),
                    "method": "easyocr",
                }, round(avg_conf, 3)
        except ImportError:
            pass
        except Exception:
            pass

        # Tesseract fallback
        try:
            import pytesseract
            from PIL import Image, ImageEnhance
            img = Image.open(file_path)
            enhanced = ImageEnhance.Contrast(img).enhance(1.3)
            text = pytesseract.image_to_string(enhanced, lang=lang)
            img.close()
            return {
                "text": text.strip(),
                "char_count": len(text.strip()),
                "method": "tesseract",
            }, 0.7
        except Exception as e:
            return {"error": f"Document OCR failed: {str(e)}", "text": ""}, None

    return await asyncio.to_thread(_doc_ocr)


async def run_document_pdf_parse(file_path: str, params: dict) -> tuple[dict, float | None]:
    """Extract text from PDF using PyMuPDF with table detection and OCR fallback."""

    def _parse():
        try:
            import fitz
        except ImportError:
            return {"error": "PyMuPDF (fitz) not installed. Install with: pip install PyMuPDF"}, None

        try:
            doc = fitz.open(file_path)
            pages = []
            full_text = []
            tables_found = []
            images_found = []
            max_pages = int(params.get("max_pages", 100))

            for page_num in range(min(doc.page_count, max_pages)):
                page = doc[page_num]
                text = page.get_text()

                # OCR fallback for scanned pages
                if len(text.strip()) < 50:
                    try:
                        import pytesseract
                        from PIL import Image as PILImage
                        pix = page.get_pixmap(dpi=200)
                        img = PILImage.frombytes("RGB", (pix.width, pix.height), pix.samples)
                        text = pytesseract.image_to_string(img)
                    except Exception:
                        pass

                pages.append({
                    "page_number": page_num + 1,
                    "text": text.strip(),
                    "char_count": len(text.strip()),
                })
                full_text.append(text.strip())

                # Extract tables
                try:
                    page_tables = page.find_tables()
                    if page_tables and page_tables.tables:
                        for table in page_tables.tables:
                            table_data = table.extract()
                            if table_data:
                                tables_found.append({
                                    "page": page_num + 1,
                                    "rows": len(table_data),
                                    "cols": len(table_data[0]) if table_data else 0,
                                    "data": table_data[:10],
                                })
                except Exception:
                    pass

                # Count embedded images
                img_list = page.get_images()
                if img_list:
                    for img_info in img_list:
                        images_found.append({
                            "page": page_num + 1,
                            "xref": img_info[0],
                            "width": img_info[2],
                            "height": img_info[3],
                        })

            metadata = doc.metadata
            total_pages_in_doc = doc.page_count
            doc.close()

            combined = "\n\n".join(full_text)
            return {
                "text": combined,
                "page_count": len(pages),
                "total_pages_in_document": total_pages_in_doc,
                "char_count": len(combined),
                "pages": pages[:10],
                "tables_found": len(tables_found),
                "tables": tables_found[:5],
                "embedded_images": len(images_found),
                "images": images_found[:20],
                "metadata": {
                    "title": metadata.get("title", ""),
                    "author": metadata.get("author", ""),
                    "subject": metadata.get("subject", ""),
                    "creator": metadata.get("creator", ""),
                    "creation_date": metadata.get("creationDate", ""),
                    "modification_date": metadata.get("modDate", ""),
                    "producer": metadata.get("producer", ""),
                    "keywords": metadata.get("keywords", ""),
                },
            }, 0.95
        except Exception as e:
            return {"error": f"PDF parsing failed: {str(e)}"}, None

    return await asyncio.to_thread(_parse)


async def run_digital_hash(file_path: str, params: dict) -> tuple[dict, float | None]:
    """Compute SHA-256, MD5, and SHA-1 hashes of a file."""

    def _hash():
        try:
            sha256 = hashlib.sha256()
            md5 = hashlib.md5()
            sha1 = hashlib.sha1()
            file_size = 0

            with open(file_path, "rb") as f:
                while True:
                    chunk = f.read(65536)
                    if not chunk:
                        break
                    sha256.update(chunk)
                    md5.update(chunk)
                    sha1.update(chunk)
                    file_size += len(chunk)

            return {
                "sha256": sha256.hexdigest(),
                "md5": md5.hexdigest(),
                "sha1": sha1.hexdigest(),
                "file_size_bytes": file_size,
                "computed_at": datetime.utcnow().isoformat(),
            }, 1.0
        except Exception as e:
            return {"error": f"Hash computation failed: {str(e)}"}, None

    return await asyncio.to_thread(_hash)


async def run_digital_metadata(file_path: str, params: dict) -> tuple[dict, float | None]:
    """Extract file size, type, creation/modification dates."""

    def _metadata():
        try:
            stat = os.stat(file_path)
            ext = os.path.splitext(file_path)[1].lower()

            import mimetypes
            mime_type, _ = mimetypes.guess_type(file_path)

            result = {
                "filename": os.path.basename(file_path),
                "extension": ext,
                "mime_type": mime_type or "application/octet-stream",
                "file_size_bytes": stat.st_size,
                "file_size_readable": _format_size(stat.st_size),
                "created_at": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                "accessed_at": datetime.fromtimestamp(stat.st_atime).isoformat(),
            }

            try:
                from PIL import Image
                if ext in (".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff", ".webp"):
                    img = Image.open(file_path)
                    result["image_width"] = img.width
                    result["image_height"] = img.height
                    result["image_mode"] = img.mode
                    img.close()
            except Exception:
                pass

            return result, 1.0
        except Exception as e:
            return {"error": f"Metadata extraction failed: {str(e)}"}, None

    return await asyncio.to_thread(_metadata)


async def run_digital_file_identify(file_path: str, params: dict) -> tuple[dict, float | None]:
    """Identify file type by magic bytes."""

    def _identify():
        try:
            with open(file_path, "rb") as f:
                header = f.read(32)

            file_signatures = {
                b"\x89PNG\r\n\x1a\n": "PNG Image",
                b"\xff\xd8\xff": "JPEG Image",
                b"GIF87a": "GIF Image (87a)",
                b"GIF89a": "GIF Image (89a)",
                b"%PDF": "PDF Document",
                b"PK\x03\x04": "ZIP Archive (or DOCX/XLSX/PPTX)",
                b"PK\x05\x06": "ZIP Archive (empty)",
                b"\x50\x4b\x07\x08": "ZIP Archive (spanned)",
                b"\x1f\x8b": "GZIP Compressed",
                b"BM": "BMP Image",
                b"\x00\x00\x01\x00": "ICO Image",
                b"RIFF": "RIFF Container (WAV/AVI)",
                b"OggS": "OGG Container",
                b"\x49\x44\x33": "MP3 Audio (ID3)",
                b"\xff\xfb": "MP3 Audio",
                b"\xff\xf3": "MP3 Audio",
                b"\xff\xf2": "MP3 Audio",
                b"\x1a\x45\xdf\xa3": "MKV/WebM Video",
                b"\x00\x00\x00": "MP4/MOV Video (likely)",
                b"\x52\x61\x72\x21": "RAR Archive",
                b"\x7f\x45\x4c\x46": "ELF Executable",
                b"\x4d\x5a": "Windows Executable (PE/DLL)",
                b"\xd0\xcf\x11\xe0": "Microsoft Office Document (OLE2)",
                b"SQLite format 3": "SQLite Database",
            }

            detected_type = "Unknown"
            for signature, file_type in file_signatures.items():
                if header.startswith(signature):
                    detected_type = file_type
                    break

            import mimetypes
            ext = os.path.splitext(file_path)[1].lower()
            mime_type, _ = mimetypes.guess_type(file_path)

            return {
                "detected_type": detected_type,
                "magic_bytes_hex": header[:16].hex(),
                "extension": ext,
                "mime_type_by_extension": mime_type or "unknown",
                "file_size_bytes": os.path.getsize(file_path),
                "header_printable": "".join(
                    chr(b) if 32 <= b < 127 else "." for b in header[:16]
                ),
            }, 0.9 if detected_type != "Unknown" else 0.3
        except Exception as e:
            return {"error": f"File identification failed: {str(e)}"}, None

    return await asyncio.to_thread(_identify)


async def run_document_summarize(file_path: str, params: dict) -> tuple[dict, float | None]:
    """Summarize document text using Gemini AI."""

    def _extract_text():
        ext = os.path.splitext(file_path)[1].lower()
        text = ""

        if ext == ".pdf":
            try:
                import fitz
                doc = fitz.open(file_path)
                for page in doc:
                    text += page.get_text() + "\n"
                doc.close()
            except Exception:
                pass
        elif ext in (".docx", ".doc"):
            try:
                import docx
                doc = docx.Document(file_path)
                text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            except Exception:
                pass
        elif ext in (".txt", ".md", ".csv", ".log", ".json", ".ipynb"):
            try:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read(50000)
            except Exception:
                pass
        elif ext in (".jpg", ".jpeg", ".png", ".tiff", ".bmp", ".webp"):
            try:
                import pytesseract
                from PIL import Image
                img = Image.open(file_path)
                text = pytesseract.image_to_string(img)
                img.close()
            except Exception:
                pass
        else:
            try:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read(50000)
            except Exception:
                pass

        return text.strip()

    text = await asyncio.to_thread(_extract_text)

    if not text:
        return {"error": "Could not extract text from the file for summarization", "summary": ""}, None

    try:
        from app.ai.llm_provider import has_any_llm_key, generate_text

        if not has_any_llm_key():
            return {
                "error": "No LLM API key configured",
                "summary": "",
                "text_preview": text[:2000],
                "original_char_count": len(text),
                "summarized_from": os.path.basename(file_path),
            }, None

        prompt = (
            "You are a forensic document analyst for a police investigation system. "
            "Provide a concise, professional summary of the following document text. "
            "Focus on key facts, names, dates, locations, and any legally relevant information.\n\n"
            f"Document text (first 10000 chars):\n{text[:10000]}\n\n"
            "Provide: 1) Brief summary (2-3 sentences) 2) Key entities found "
            "3) Potential investigative relevance"
        )

        summary = await asyncio.to_thread(generate_text, prompt, 0.3, 2048)

        return {
            "summary": summary,
            "original_char_count": len(text),
            "summarized_from": os.path.basename(file_path),
        }, 0.8
    except ImportError:
        return {
            "error": "Required LLM packages not installed",
            "summary": "",
            "text_preview": text[:2000],
        }, None
    except Exception as e:
        return {
            "error": f"AI summarization failed: {str(e)}",
            "summary": "",
            "text_preview": text[:2000],
        }, None


# ---------------------------------------------------------------------------
# Face Detection (OpenCV Haar Cascade + InsightFace)
# ---------------------------------------------------------------------------

async def run_face_detect(file_path: str, params: dict) -> tuple[dict, float | None]:
    """Detect faces in image using InsightFace (preferred) or OpenCV fallback."""

    def _detect_faces():
        try:
            import cv2
            img = cv2.imread(file_path)
            if img is None:
                return {"error": "Could not read image file"}, None
        except ImportError:
            return {"error": "OpenCV not available"}, None

        # Try InsightFace first
        try:
            app = _get_insightface()
            faces = app.get(img)

            face_list = []
            for i, face in enumerate(faces):
                bbox = face.bbox.astype(int).tolist()
                face_list.append({
                    "face_id": i + 1,
                    "bbox": {
                        "x": bbox[0], "y": bbox[1],
                        "width": bbox[2] - bbox[0], "height": bbox[3] - bbox[1]
                    },
                    "confidence": round(float(face.det_score), 4),
                    "age": int(face.age) if hasattr(face, "age") and face.age else None,
                    "gender": "male" if hasattr(face, "gender") and face.gender == 1 else "female" if hasattr(face, "gender") and face.gender == 0 else None,
                })

            return {
                "faces_detected": len(face_list),
                "faces": face_list,
                "image_dimensions": {"width": img.shape[1], "height": img.shape[0]},
                "method": "InsightFace (buffalo_l)",
            }, round(sum(f["confidence"] for f in face_list) / len(face_list), 3) if face_list else None

        except Exception as e:
            logger.warning(f"InsightFace failed, falling back to Haar: {e}")

        # Fallback to Haar Cascade
        try:
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            cascade_paths = [
                cv2.data.haarcascades + "haarcascade_frontalface_default.xml",
                cv2.data.haarcascades + "haarcascade_frontalface_alt2.xml",
            ]

            face_cascade = None
            for path in cascade_paths:
                if os.path.exists(path):
                    face_cascade = cv2.CascadeClassifier(path)
                    break

            if face_cascade is None:
                return {"error": "No face detection model available", "faces": []}, None

            faces = face_cascade.detectMultiScale(gray, scaleFactor=1.1, minNeighbors=5, minSize=(30, 30))

            face_list = []
            for i, (x, y, w, h) in enumerate(faces):
                face_list.append({
                    "face_id": i + 1,
                    "bbox": {"x": int(x), "y": int(y), "width": int(w), "height": int(h)},
                })

            return {
                "faces_detected": len(face_list),
                "faces": face_list,
                "image_dimensions": {"width": img.shape[1], "height": img.shape[0]},
                "method": "Haar Cascade (OpenCV)",
            }, 0.7 if face_list else None

        except Exception as e:
            return {"error": f"Face detection failed: {str(e)}", "faces": []}, None

    return await asyncio.to_thread(_detect_faces)


# ---------------------------------------------------------------------------
# Face Recognition (InsightFace embeddings + Qdrant/DB search)
# ---------------------------------------------------------------------------

async def run_face_recognize(file_path: str, params: dict) -> tuple[dict, float | None]:
    """Detect faces, generate embeddings via InsightFace, search criminal DB."""

    def _recognize():
        start_time = time.time()

        try:
            import cv2
            img = cv2.imread(file_path)
            if img is None:
                return {"error": "Could not read image file"}, None
        except ImportError:
            return {"error": "OpenCV not available"}, None

        try:
            app = _get_insightface()
        except Exception as e:
            return {"error": f"InsightFace model failed to load: {str(e)}"}, None

        faces = app.get(img)
        if not faces:
            return {
                "faces_detected": 0,
                "matches": [],
                "image_dimensions": {"width": img.shape[1], "height": img.shape[0]},
                "method": "InsightFace (buffalo_l)",
                "database_searched": "criminal_face_embeddings",
                "database_records_searched": 0,
                "note": "No faces detected in the uploaded image.",
            }, None

        threshold = float(params.get("threshold", 0.4))
        top_k = int(params.get("top_k", 5))

        # Try Qdrant first, then fall back to DB direct comparison
        all_matches = []
        db_records_searched = 0

        try:
            client = _get_qdrant_client()
            _ensure_qdrant_collection(client, FACE_COLLECTION, FACE_EMBEDDING_DIM)

            collection_info = client.get_collection(FACE_COLLECTION)
            db_records_searched = collection_info.points_count

            if db_records_searched > 0:
                for face_idx, face in enumerate(faces):
                    embedding = face.normed_embedding.tolist()
                    search_results = client.search(
                        collection_name=FACE_COLLECTION,
                        query_vector=embedding,
                        limit=top_k,
                        score_threshold=threshold,
                    )

                    for result in search_results:
                        payload = result.payload or {}
                        all_matches.append({
                            "face_index": face_idx + 1,
                            "criminal_id": payload.get("criminal_id", ""),
                            "criminal_name": payload.get("full_name", "Unknown"),
                            "similarity": round(result.score, 4),
                            "wanted_status": payload.get("wanted_status", ""),
                            "danger_level": payload.get("danger_level", ""),
                            "gang_name": payload.get("gang_name"),
                            "source_image": payload.get("image_path", ""),
                        })
        except Exception as e:
            logger.warning(f"Qdrant face search failed, falling back to DB: {e}")

        # Fallback: direct comparison against PostgreSQL stored embeddings
        if db_records_searched == 0:
            try:
                db_embeddings = _query_all_face_embeddings_sync()
                db_records_searched = len(db_embeddings)

                for face_idx, face in enumerate(faces):
                    query_emb = face.normed_embedding
                    face_matches = []

                    for record in db_embeddings:
                        stored_emb = np.array(record["embedding"], dtype=np.float32)
                        if stored_emb.shape[0] != FACE_EMBEDDING_DIM:
                            continue
                        stored_emb = stored_emb / (np.linalg.norm(stored_emb) + 1e-8)
                        similarity = float(np.dot(query_emb, stored_emb))

                        if similarity >= threshold:
                            face_matches.append({
                                "face_index": face_idx + 1,
                                "criminal_id": record["criminal_id"],
                                "criminal_name": record["full_name"],
                                "similarity": round(similarity, 4),
                                "wanted_status": record["wanted_status"],
                                "danger_level": record["danger_level"],
                                "gang_name": record["gang_name"],
                                "source_image": record.get("image_path", ""),
                            })

                    face_matches.sort(key=lambda x: x["similarity"], reverse=True)
                    all_matches.extend(face_matches[:top_k])
            except Exception as e:
                logger.error(f"DB face comparison failed: {e}")

        # Sort all matches by similarity
        all_matches.sort(key=lambda x: x["similarity"], reverse=True)
        all_matches = all_matches[:top_k]

        # Build face info
        face_details = []
        for i, face in enumerate(faces):
            bbox = face.bbox.astype(int).tolist()
            face_details.append({
                "face_id": i + 1,
                "bbox": {
                    "x": bbox[0], "y": bbox[1],
                    "width": bbox[2] - bbox[0], "height": bbox[3] - bbox[1]
                },
                "confidence": round(float(face.det_score), 4),
                "embedding_generated": True,
                "embedding_dim": FACE_EMBEDDING_DIM,
            })

        elapsed = round((time.time() - start_time) * 1000, 1)
        top_confidence = all_matches[0]["similarity"] if all_matches else None

        return {
            "faces_detected": len(faces),
            "faces": face_details,
            "matches_found": len(all_matches),
            "matches": all_matches,
            "image_dimensions": {"width": img.shape[1], "height": img.shape[0]},
            "method": "InsightFace (buffalo_l) + Cosine Similarity",
            "database_searched": "criminal_face_embeddings",
            "database_records_searched": db_records_searched,
            "threshold_used": threshold,
            "execution_time_ms": elapsed,
        }, top_confidence

    output_data, confidence = await asyncio.to_thread(_recognize)
    if output_data.get("matches"):
        output_data["matches"] = await _enrich_matches_with_profiles(output_data["matches"])
    return output_data, confidence


# ---------------------------------------------------------------------------
# Fingerprint Matching (OpenCV ORB feature extraction + template matching)
# ---------------------------------------------------------------------------

async def run_fingerprint_match(file_path: str, params: dict) -> tuple[dict, float | None]:
    """Extract fingerprint features using OpenCV and match against criminal DB."""

    def _match():
        start_time = time.time()

        try:
            import cv2
        except ImportError:
            return {"error": "OpenCV not available for fingerprint processing"}, None

        img = cv2.imread(file_path, cv2.IMREAD_GRAYSCALE)
        if img is None:
            return {"error": "Could not read fingerprint image"}, None

        # Fingerprint preprocessing
        # 1. Histogram equalization for contrast
        equalized = cv2.equalizeHist(img)

        # 2. Gaussian blur to reduce noise
        blurred = cv2.GaussianBlur(equalized, (5, 5), 0)

        # 3. Adaptive thresholding for ridge enhancement
        binary = cv2.adaptiveThreshold(
            blurred, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY_INV, 11, 2
        )

        # 4. Morphological operations for ridge cleanup
        kernel = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (3, 3))
        cleaned = cv2.morphologyEx(binary, cv2.MORPH_CLOSE, kernel)

        # 5. Thin ridges using skeletonization approximation
        thinned = cv2.ximgproc.thinning(cleaned) if hasattr(cv2, 'ximgproc') else cleaned

        # 6. Extract minutiae using ORB (ridge endings and bifurcations approximation)
        orb = cv2.ORB_create(nfeatures=500, scaleFactor=1.2, nlevels=8)
        keypoints, descriptors = orb.detectAndCompute(thinned, None)

        if descriptors is None or len(keypoints) < 10:
            # Try SIFT as fallback
            try:
                sift = cv2.SIFT_create(nfeatures=500)
                keypoints, descriptors = sift.detectAndCompute(img, None)
            except Exception:
                pass

        if descriptors is None or len(keypoints) < 5:
            return {
                "error": "Insufficient fingerprint features detected. Image quality too low.",
                "minutiae_count": len(keypoints) if keypoints else 0,
                "quality_assessment": "poor",
            }, None

        # Build template
        input_template = {
            "minutiae_count": len(keypoints),
            "descriptors": descriptors.tolist()[:200],
            "keypoints": [
                {"x": round(kp.pt[0], 1), "y": round(kp.pt[1], 1),
                 "size": round(kp.size, 1), "angle": round(kp.angle, 1)}
                for kp in keypoints[:200]
            ],
        }

        # Quality assessment
        minutiae_count = len(keypoints)
        if minutiae_count >= 100:
            quality = "excellent"
            quality_score = 0.95
        elif minutiae_count >= 50:
            quality = "good"
            quality_score = 0.8
        elif minutiae_count >= 20:
            quality = "fair"
            quality_score = 0.6
        else:
            quality = "poor"
            quality_score = 0.3

        # Search criminal fingerprint database
        threshold = float(params.get("threshold", 0.35))
        top_k = int(params.get("top_k", 5))
        matches = []
        db_records_searched = 0

        try:
            db_fingerprints = _query_all_fingerprints_sync()
            db_records_searched = len(db_fingerprints)

            bf = cv2.BFMatcher(cv2.NORM_HAMMING, crossCheck=False)

            for record in db_fingerprints:
                template = record.get("template_data")
                if not template or "descriptors" not in template:
                    continue

                try:
                    stored_desc = np.array(template["descriptors"], dtype=np.uint8)
                    if stored_desc.shape[0] < 5:
                        continue

                    raw_matches = bf.knnMatch(descriptors, stored_desc, k=2)

                    # Lowe's ratio test
                    good_matches = []
                    for m_pair in raw_matches:
                        if len(m_pair) == 2:
                            m, n = m_pair
                            if m.distance < 0.75 * n.distance:
                                good_matches.append(m)

                    # Calculate similarity score
                    max_possible = min(len(descriptors), len(stored_desc))
                    similarity = len(good_matches) / max_possible if max_possible > 0 else 0

                    if similarity >= threshold:
                        matches.append({
                            "criminal_id": record["criminal_id"],
                            "criminal_name": record["full_name"],
                            "finger_type": record["finger_type"],
                            "similarity": round(similarity, 4),
                            "matching_points": len(good_matches),
                            "total_points_compared": max_possible,
                            "wanted_status": record["wanted_status"],
                            "danger_level": record["danger_level"],
                        })
                except Exception:
                    continue

            matches.sort(key=lambda x: x["similarity"], reverse=True)
            matches = matches[:top_k]

        except Exception as e:
            logger.error(f"Fingerprint DB search failed: {e}")

        elapsed = round((time.time() - start_time) * 1000, 1)
        top_confidence = matches[0]["similarity"] if matches else None

        return {
            "minutiae_extracted": minutiae_count,
            "quality_assessment": quality,
            "quality_score": quality_score,
            "template_generated": True,
            "matches_found": len(matches),
            "matches": matches,
            "database_searched": "criminal_fingerprints",
            "database_records_searched": db_records_searched,
            "threshold_used": threshold,
            "method": "ORB Feature Extraction + BFMatcher + Lowe's Ratio Test",
            "execution_time_ms": elapsed,
            "input_template_summary": {
                "minutiae_count": minutiae_count,
                "keypoints_sample": input_template["keypoints"][:5],
            },
        }, top_confidence

    output_data, confidence = await asyncio.to_thread(_match)
    if output_data.get("matches"):
        output_data["matches"] = await _enrich_matches_with_profiles(output_data["matches"])
    return output_data, confidence


# ---------------------------------------------------------------------------
# DNA Analysis (OCR + Profile Parsing + DB Search)
# ---------------------------------------------------------------------------

async def run_dna_search(file_path: str, params: dict) -> tuple[dict, float | None]:
    """Parse DNA report via OCR, extract profile, and search criminal DB."""

    def _search():
        start_time = time.time()

        # Step 1: Extract text from file
        ext = os.path.splitext(file_path)[1].lower()
        text_content = ""

        if ext in (".txt", ".csv", ".tsv", ".json"):
            try:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    text_content = f.read(50000)
            except Exception:
                pass
        elif ext == ".pdf":
            try:
                import fitz
                doc = fitz.open(file_path)
                for page in doc:
                    text_content += page.get_text() + "\n"
                doc.close()
            except Exception:
                pass
        else:
            try:
                import pytesseract
                from PIL import Image
                img = Image.open(file_path)
                text_content = pytesseract.image_to_string(img)
                img.close()
            except Exception:
                pass

        if not text_content.strip():
            return {
                "error": "Could not extract text from the DNA report file.",
                "matches": [],
            }, None

        # Step 2: Parse DNA profile data
        parsed_profile = _parse_dna_report(text_content)

        # Step 3: Search criminal DNA profiles
        threshold = float(params.get("threshold", 0.7))
        matches = []
        db_records_searched = 0

        try:
            db_profiles = _query_all_dna_profiles_sync()
            db_records_searched = len(db_profiles)

            input_loci = parsed_profile.get("loci", {})

            for record in db_profiles:
                stored_loci = record.get("loci_markers")
                if not stored_loci or not isinstance(stored_loci, dict):
                    continue

                # Compare STR loci
                similarity = _compare_dna_loci(input_loci, stored_loci)

                if similarity >= threshold:
                    matches.append({
                        "criminal_id": record["criminal_id"],
                        "criminal_name": record["full_name"],
                        "dna_id": record["dna_id"],
                        "similarity": round(similarity, 4),
                        "matching_loci": _count_matching_loci(input_loci, stored_loci),
                        "total_loci_compared": len(set(input_loci.keys()) & set(stored_loci.keys())),
                        "wanted_status": record["wanted_status"],
                        "danger_level": record["danger_level"],
                        "laboratory": record.get("laboratory"),
                    })

            matches.sort(key=lambda x: x["similarity"], reverse=True)
            matches = matches[:5]

        except Exception as e:
            logger.error(f"DNA DB search failed: {e}")

        elapsed = round((time.time() - start_time) * 1000, 1)
        top_confidence = matches[0]["similarity"] if matches else None

        return {
            "report_parsed": True,
            "extracted_profile": parsed_profile,
            "text_preview": text_content[:1000],
            "matches_found": len(matches),
            "matches": matches,
            "database_searched": "criminal_dna_profiles",
            "database_records_searched": db_records_searched,
            "threshold_used": threshold,
            "method": "OCR + STR Loci Parsing + Allele Comparison",
            "execution_time_ms": elapsed,
        }, top_confidence

    output_data, confidence = await asyncio.to_thread(_search)
    if output_data.get("matches"):
        output_data["matches"] = await _enrich_matches_with_profiles(output_data["matches"])
    return output_data, confidence


def _parse_dna_report(text: str) -> dict:
    """Parse DNA report text to extract profile information."""
    profile = {
        "dna_id": None,
        "sample_number": None,
        "laboratory": None,
        "collection_date": None,
        "officer": None,
        "case_reference": None,
        "loci": {},
    }

    # Extract DNA Profile ID
    id_patterns = [
        r"(?:DNA\s*(?:Profile\s*)?ID|Profile\s*(?:No|Number|ID))[:\s]*([A-Z0-9\-/]+)",
        r"(?:Sample\s*ID|Specimen\s*ID)[:\s]*([A-Z0-9\-/]+)",
    ]
    for pat in id_patterns:
        match = re.search(pat, text, re.IGNORECASE)
        if match:
            profile["dna_id"] = match.group(1).strip()
            break

    # Extract sample number
    sample_match = re.search(r"(?:Sample\s*(?:No|Number)|Exhibit\s*No)[:\s]*([A-Z0-9\-/]+)", text, re.IGNORECASE)
    if sample_match:
        profile["sample_number"] = sample_match.group(1).strip()

    # Extract laboratory
    lab_match = re.search(r"(?:Laboratory|Lab|FSL|Forensic\s*Science)[:\s]*([^\n,]+)", text, re.IGNORECASE)
    if lab_match:
        profile["laboratory"] = lab_match.group(1).strip()[:100]

    # Extract date
    date_match = re.search(r"(?:Date|Collection\s*Date|Report\s*Date)[:\s]*(\d{1,2}[\-/]\d{1,2}[\-/]\d{2,4})", text, re.IGNORECASE)
    if date_match:
        profile["collection_date"] = date_match.group(1).strip()

    # Extract officer
    officer_match = re.search(r"(?:Officer|Investigating\s*Officer|IO)[:\s]*([A-Za-z\s\.]+)", text, re.IGNORECASE)
    if officer_match:
        profile["officer"] = officer_match.group(1).strip()[:100]

    # Extract case reference
    case_match = re.search(r"(?:Case|FIR|Crime)\s*(?:No|Number|Ref)[:\s]*([A-Z0-9\-/]+)", text, re.IGNORECASE)
    if case_match:
        profile["case_reference"] = case_match.group(1).strip()

    # Extract STR loci
    standard_loci = [
        "D3S1358", "vWA", "D16S539", "CSF1PO", "TPOX", "D8S1179",
        "D21S11", "D18S51", "D2S441", "D19S433", "TH01", "FGA",
        "D22S1045", "D5S818", "D13S317", "D7S820", "SE33", "D10S1248",
        "D1S1656", "D12S391", "D2S1338", "Penta D", "Penta E",
        "AMEL", "Amelogenin", "DYS391",
    ]

    for locus in standard_loci:
        pattern = rf"{re.escape(locus)}\s*[:\|\t,]\s*(\d+(?:\.\d+)?)\s*[/,\|\s]\s*(\d+(?:\.\d+)?)"
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            profile["loci"][locus] = [float(match.group(1)), float(match.group(2))]

    # Also try tab/pipe-separated table format
    if not profile["loci"]:
        lines = text.split("\n")
        for line in lines:
            for locus in standard_loci:
                if locus.lower() in line.lower():
                    numbers = re.findall(r"\d+(?:\.\d+)?", line)
                    if len(numbers) >= 2:
                        profile["loci"][locus] = [float(numbers[-2]), float(numbers[-1])]
                        break

    return profile


def _compare_dna_loci(input_loci: dict, stored_loci: dict) -> float:
    """Compare two DNA profiles by their STR loci alleles."""
    if not input_loci or not stored_loci:
        return 0.0

    common_loci = set(input_loci.keys()) & set(stored_loci.keys())
    if not common_loci:
        return 0.0

    matches = 0
    total = len(common_loci) * 2  # 2 alleles per locus

    for locus in common_loci:
        input_alleles = sorted(input_loci[locus])
        stored_alleles = sorted(stored_loci[locus])

        if len(input_alleles) >= 2 and len(stored_alleles) >= 2:
            if abs(input_alleles[0] - stored_alleles[0]) < 0.1:
                matches += 1
            if abs(input_alleles[1] - stored_alleles[1]) < 0.1:
                matches += 1

    return matches / total if total > 0 else 0.0


def _count_matching_loci(input_loci: dict, stored_loci: dict) -> int:
    """Count fully matching loci between two profiles."""
    count = 0
    common = set(input_loci.keys()) & set(stored_loci.keys())
    for locus in common:
        input_alleles = sorted(input_loci[locus])
        stored_alleles = sorted(stored_loci[locus])
        if len(input_alleles) >= 2 and len(stored_alleles) >= 2:
            if (abs(input_alleles[0] - stored_alleles[0]) < 0.1 and
                    abs(input_alleles[1] - stored_alleles[1]) < 0.1):
                count += 1
    return count


# ---------------------------------------------------------------------------
# Vehicle Detection
# ---------------------------------------------------------------------------

async def run_vehicle_detect(file_path: str, params: dict) -> tuple[dict, float | None]:
    """Detect vehicles in image using YOLO with Gemini Vision fallback."""

    def _detect_yolo():
        try:
            from ultralytics import YOLO
        except ImportError:
            return None

        try:
            vehicle_classes = {2: "car", 3: "motorcycle", 5: "bus", 7: "truck", 1: "bicycle"}

            model_path = params.get("model", "yolov8n.pt")
            confidence_threshold = float(params.get("confidence", 0.15))
            model = YOLO(model_path)
            results = model(file_path, conf=confidence_threshold)

            vehicles = []
            total_conf = 0.0

            for result in results:
                for box in result.boxes:
                    cls_id = int(box.cls[0])
                    if cls_id in vehicle_classes:
                        conf = float(box.conf[0])
                        x1, y1, x2, y2 = box.xyxy[0].tolist()
                        vehicles.append({
                            "type": vehicle_classes[cls_id],
                            "confidence": round(conf, 4),
                            "bbox": {
                                "x1": round(x1, 1),
                                "y1": round(y1, 1),
                                "x2": round(x2, 1),
                                "y2": round(y2, 1),
                            },
                        })
                        total_conf += conf

            if vehicles:
                avg_conf = round(total_conf / len(vehicles), 3)
                return {
                    "vehicles_detected": len(vehicles),
                    "vehicles": vehicles,
                    "model_used": model_path,
                    "method": "yolo",
                }, avg_conf
            return None
        except Exception:
            return None

    def _detect_gemini():
        from app.ai.llm_provider import has_any_llm_key, generate_vision
        if not has_any_llm_key():
            return None
        try:
            prompt = (
                "Analyze this image for vehicles. For each vehicle visible, identify:\n"
                "1. Type (car, motorcycle, bus, truck, auto-rickshaw, van, SUV, bicycle)\n"
                "2. Color\n"
                "3. Approximate make/model if identifiable\n"
                "Respond in this exact format per vehicle (one per line): VEHICLE: <type> | <color> | <make_model>\n"
                "If no vehicle is visible, respond: NO_VEHICLE_FOUND\n"
                "Do not add any other text."
            )
            response_text = generate_vision(file_path, prompt, temperature=0.1, max_tokens=300)

            text = response_text.strip()
            if "NO_VEHICLE_FOUND" in text:
                return {"vehicles_detected": 0, "vehicles": [], "method": "gemini_vision"}, None

            vehicles = []
            for line in text.split("\n"):
                line = line.strip()
                if line.startswith("VEHICLE:"):
                    parts = [p.strip() for p in line[8:].split("|")]
                    v_type = parts[0] if parts else "unknown"
                    color = parts[1] if len(parts) > 1 else "unknown"
                    make = parts[2] if len(parts) > 2 else "unknown"
                    vehicles.append({
                        "type": v_type.lower(),
                        "color": color,
                        "make_model": make,
                        "confidence": 0.9,
                        "method": "gemini_vision",
                    })

            return {
                "vehicles_detected": len(vehicles),
                "vehicles": vehicles,
                "method": "gemini_vision",
            }, 0.9 if vehicles else None
        except Exception:
            return None

    def _run():
        result = _detect_yolo()
        if result is not None:
            return result
        result = _detect_gemini()
        if result is not None:
            return result
        return {"vehicles_detected": 0, "vehicles": [], "method": "yolo (no detection)"}, None

    return await asyncio.to_thread(_run)


async def run_license_plate_ocr(file_path: str, params: dict) -> tuple[dict, float | None]:
    """Detect license plate using Gemini Vision (primary) with Tesseract fallback."""

    def _plate_ocr_gemini():
        """Use Vision API to read license plates — far more accurate than Tesseract on natural photos."""
        from app.ai.llm_provider import has_any_llm_key, generate_vision
        if not has_any_llm_key():
            return None

        try:
            prompt = (
                "You are a license plate detection system. Look at this image and find ALL vehicle license plates visible. "
                "For each plate found, extract the exact text on it. "
                "Respond ONLY in this exact format (one plate per line): PLATE: <plate_text>\n"
                "If no license plate is visible, respond with exactly: NO_PLATE_FOUND\n"
                "Do not add any other text, explanation, or commentary."
            )
            response_text = generate_vision(file_path, prompt, temperature=0.1, max_tokens=200).strip()

            if "NO_PLATE_FOUND" in response_text:
                return {"plates_detected": 0, "plates": [], "method": "gemini_vision"}, None

            plates = []
            for line in response_text.split("\n"):
                line = line.strip()
                if line.startswith("PLATE:"):
                    plate_text = line[6:].strip().upper()
                    plate_text = re.sub(r'[^A-Z0-9\s\-]', '', plate_text).strip()
                    if len(plate_text) >= 3:
                        plates.append({
                            "text": plate_text,
                            "method": "gemini_vision",
                            "confidence": 0.95,
                        })

            if not plates and "PLATE" not in response_text.upper():
                cleaned = re.sub(r'[^A-Z0-9\s\-]', '', response_text.upper()).strip()
                if len(cleaned) >= 4:
                    plates.append({"text": cleaned, "method": "gemini_vision_raw", "confidence": 0.7})

            return {
                "plates_detected": len(plates),
                "plates": plates,
                "method": "gemini_vision",
            }, 0.95 if plates else None

        except Exception:
            return None

    def _plate_ocr_easyocr():
        """Use EasyOCR (deep learning) — works on real-world photos unlike Tesseract."""
        try:
            import easyocr
            import cv2
            import numpy as np
        except ImportError:
            return None

        try:
            img = cv2.imread(file_path)
            if img is None:
                return None

            h_img, w_img = img.shape[:2]
            reader = easyocr.Reader(['en'], gpu=False, verbose=False)
            results = reader.readtext(file_path)

            plates = []
            for (bbox, text, conf) in results:
                text_clean = re.sub(r'[^A-Za-z0-9\s\-]', '', text).upper().strip()
                if len(text_clean) < 4:
                    continue
                (tl, tr, br, bl) = bbox
                x = int(min(tl[0], bl[0]))
                y = int(min(tl[1], tr[1]))
                w = int(max(tr[0], br[0]) - x)
                h = int(max(bl[1], br[1]) - y)
                aspect_ratio = w / float(h) if h > 0 else 0

                is_plate_like = (
                    1.2 <= aspect_ratio <= 8.0
                    and bool(re.search(r'\d', text_clean))
                    and bool(re.search(r'[A-Z]', text_clean))
                    and len(text_clean) >= 4
                )
                if is_plate_like:
                    plates.append({
                        "text": text_clean,
                        "bbox": {"x": x, "y": y, "width": w, "height": h},
                        "method": "easyocr",
                        "confidence": round(conf, 2),
                    })

            if not plates:
                for (bbox, text, conf) in results:
                    text_clean = re.sub(r'[^A-Za-z0-9\s\-]', '', text).upper().strip()
                    if len(text_clean) >= 5 and conf > 0.3:
                        has_mixed = bool(re.search(r'\d', text_clean)) and bool(re.search(r'[A-Z]', text_clean))
                        if has_mixed:
                            (tl, tr, br, bl) = bbox
                            plates.append({
                                "text": text_clean,
                                "bbox": {"x": int(tl[0]), "y": int(tl[1]),
                                         "width": int(tr[0] - tl[0]), "height": int(bl[1] - tl[1])},
                                "method": "easyocr_relaxed",
                                "confidence": round(conf, 2),
                            })

            if plates:
                return {
                    "plates_detected": len(plates),
                    "plates": plates[:5],
                    "method": "easyocr_deep_learning",
                    "image_dimensions": {"width": w_img, "height": h_img},
                    "total_text_regions": len(results),
                }, 0.85 if plates else None

            return None
        except Exception:
            return None

    def _plate_ocr_tesseract():
        """Last-resort fallback: Tesseract (works on clean/synthetic images only)."""
        try:
            from PIL import Image, ImageEnhance
            import pytesseract
            import cv2
        except ImportError as e:
            return {"error": f"Required library not available: {e}", "plates": []}, None

        try:
            img = cv2.imread(file_path)
            if img is None:
                return {"error": "Could not read image"}, None

            h_img, w_img = img.shape[:2]
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            plates = []

            pil_img = Image.open(file_path)
            enhanced = ImageEnhance.Contrast(pil_img).enhance(1.5)
            enhanced = ImageEnhance.Sharpness(enhanced).enhance(2.0)
            full_text = pytesseract.image_to_string(enhanced, config="--psm 6 --oem 3")
            pil_img.close()
            plate_patterns = re.findall(
                r"[A-Z]{2}[\s\-]?\d{1,2}[\s\-]?[A-Z]{1,3}[\s\-]?\d{1,4}", full_text.upper()
            )
            if not plate_patterns:
                plate_patterns = re.findall(
                    r"[A-Z0-9]{2,4}[\s\-]?[A-Z0-9]{2,4}[\s\-]?[A-Z0-9]{1,4}", full_text.upper()
                )
                plate_patterns = [p for p in plate_patterns if len(p.replace(' ', '').replace('-', '')) >= 5]
            for pat in plate_patterns[:5]:
                plates.append({"text": pat.strip(), "method": "tesseract_fallback", "confidence": 0.4})

            return {
                "plates_detected": len(plates),
                "plates": plates,
                "method": "tesseract_fallback",
                "image_dimensions": {"width": w_img, "height": h_img},
            }, 0.5 if plates else None

        except Exception as e:
            return {"error": f"License plate OCR failed: {str(e)}", "plates": []}, None

    def _run():
        result = _plate_ocr_gemini()
        if result is not None:
            return result
        result = _plate_ocr_easyocr()
        if result is not None:
            return result
        return _plate_ocr_tesseract()

    return await asyncio.to_thread(_run)


async def run_weapon_detect(file_path: str, params: dict) -> tuple[dict, float | None]:
    """Detect weapons in image using YOLO with Gemini Vision fallback."""

    def _detect_yolo():
        try:
            from ultralytics import YOLO
        except ImportError:
            return None

        try:
            model_path = params.get("model", "yolov8n.pt")
            confidence_threshold = float(params.get("confidence", 0.25))
            model = YOLO(model_path)
            results = model(file_path, conf=confidence_threshold)

            weapon_keywords = {"knife", "scissors", "baseball bat", "gun", "rifle", "pistol", "sword"}
            weapons = []
            all_objects = []
            total_conf = 0.0

            for result in results:
                for box in result.boxes:
                    cls_id = int(box.cls[0])
                    cls_name = result.names[cls_id].lower()
                    conf = float(box.conf[0])
                    x1, y1, x2, y2 = box.xyxy[0].tolist()

                    obj_data = {
                        "class": cls_name,
                        "confidence": round(conf, 4),
                        "bbox": {
                            "x1": round(x1, 1),
                            "y1": round(y1, 1),
                            "x2": round(x2, 1),
                            "y2": round(y2, 1),
                        },
                    }

                    all_objects.append(obj_data)
                    if cls_name in weapon_keywords:
                        weapons.append(obj_data)
                        total_conf += conf

            avg_conf = round(total_conf / len(weapons), 3) if weapons else None
            return {
                "weapons_detected": len(weapons),
                "weapons": weapons,
                "all_objects_detected": len(all_objects),
                "model_used": model_path,
                "method": "yolo",
            }, avg_conf
        except Exception:
            return None

    def _detect_gemini():
        from app.ai.llm_provider import has_any_llm_key, generate_vision
        if not has_any_llm_key():
            return None
        try:
            prompt = (
                "You are a weapon detection system for law enforcement. Analyze this image carefully.\n"
                "Identify ANY weapons or potentially dangerous objects: guns, pistols, rifles, shotguns, "
                "knives, machetes, swords, axes, baseball bats, brass knuckles, explosives, etc.\n"
                "For each weapon found respond: WEAPON: <type> | <description>\n"
                "If NO weapons are visible, respond exactly: NO_WEAPON_FOUND\n"
                "Do not add any other text. Be accurate — do not hallucinate weapons that aren't there."
            )
            text = generate_vision(file_path, prompt, temperature=0.1, max_tokens=300).strip()

            if "NO_WEAPON_FOUND" in text:
                return {
                    "weapons_detected": 0,
                    "weapons": [],
                    "method": "gemini_vision",
                    "threat_level": "none",
                }, None

            weapons = []
            for line in text.split("\n"):
                line = line.strip()
                if line.startswith("WEAPON:"):
                    parts = [p.strip() for p in line[7:].split("|")]
                    w_type = parts[0] if parts else "unknown"
                    desc = parts[1] if len(parts) > 1 else ""
                    weapons.append({
                        "class": w_type.lower(),
                        "description": desc,
                        "confidence": 0.85,
                        "method": "gemini_vision",
                    })

            threat = "high" if weapons else "none"
            return {
                "weapons_detected": len(weapons),
                "weapons": weapons,
                "method": "gemini_vision",
                "threat_level": threat,
            }, 0.85 if weapons else None
        except Exception:
            return None

    def _run():
        yolo_result = _detect_yolo()

        if yolo_result and yolo_result[0].get("weapons_detected", 0) > 0:
            return yolo_result

        gemini_result = _detect_gemini()
        if gemini_result is not None:
            if yolo_result:
                gemini_result[0]["all_objects_detected"] = yolo_result[0].get("all_objects_detected", 0)
            return gemini_result

        if yolo_result is not None:
            return yolo_result
        return {"weapons_detected": 0, "weapons": [], "method": "none"}, None

    return await asyncio.to_thread(_run)


# ---------------------------------------------------------------------------
# Image Similarity (CLIP embeddings + Qdrant)
# ---------------------------------------------------------------------------

async def run_image_similarity(file_path: str, params: dict) -> tuple[dict, float | None]:
    """Generate CLIP embedding and search for similar images in Qdrant."""

    def _compare():
        start_time = time.time()

        try:
            import torch
            from PIL import Image
        except ImportError as e:
            return {"error": f"Required libraries not available: {e}"}, None

        try:
            img = Image.open(file_path).convert("RGB")
        except Exception as e:
            return {"error": f"Could not open image: {e}"}, None

        # Generate CLIP embedding
        try:
            model, preprocess = _get_clip_model()

            image_input = preprocess(img).unsqueeze(0)
            with torch.no_grad():
                embedding = model.encode_image(image_input)
                embedding = embedding / embedding.norm(dim=-1, keepdim=True)
                embedding_list = embedding[0].cpu().numpy().tolist()

        except Exception as e:
            return {"error": f"CLIP embedding generation failed: {e}"}, None

        # Search Qdrant for similar images
        top_k = int(params.get("top_k", 10))
        threshold = float(params.get("threshold", 0.5))
        matches = []
        db_records_searched = 0

        try:
            client = _get_qdrant_client()
            _ensure_qdrant_collection(client, IMAGE_COLLECTION, CLIP_EMBEDDING_DIM)

            collection_info = client.get_collection(IMAGE_COLLECTION)
            db_records_searched = collection_info.points_count

            if db_records_searched > 0:
                search_results = client.search(
                    collection_name=IMAGE_COLLECTION,
                    query_vector=embedding_list,
                    limit=top_k,
                    score_threshold=threshold,
                )

                for result in search_results:
                    payload = result.payload or {}
                    matches.append({
                        "image_path": payload.get("image_path", ""),
                        "filename": payload.get("filename", ""),
                        "similarity": round(result.score, 4),
                        "case_id": payload.get("case_id"),
                        "evidence_id": payload.get("evidence_id"),
                        "tool_execution_id": payload.get("execution_id"),
                        "uploaded_at": payload.get("uploaded_at", ""),
                    })

        except Exception as e:
            logger.warning(f"Qdrant image search failed: {e}")

        elapsed = round((time.time() - start_time) * 1000, 1)
        top_confidence = matches[0]["similarity"] if matches else None

        return {
            "source_image": os.path.basename(file_path),
            "source_dimensions": {"width": img.width, "height": img.height},
            "embedding_generated": True,
            "embedding_dim": CLIP_EMBEDDING_DIM,
            "embedding_model": "ViT-B-32 (OpenCLIP, laion2b_s34b_b79k)",
            "matches_found": len(matches),
            "matches": matches,
            "database_searched": "image_embeddings",
            "database_records_searched": db_records_searched,
            "threshold_used": threshold,
            "method": "CLIP ViT-B-32 + Cosine Similarity",
            "execution_time_ms": elapsed,
        }, top_confidence

    return await asyncio.to_thread(_compare)


# ---------------------------------------------------------------------------
# Crime Scene Analysis (Multimodal Pipeline)
# ---------------------------------------------------------------------------

async def run_crime_scene_analysis(file_path: str, params: dict) -> tuple[dict, float | None]:
    """Full multimodal crime scene analysis pipeline."""

    # Run all detection tools in parallel
    ocr_result, _ = await run_image_ocr(file_path, {"language": "eng"})
    object_result, _ = await run_object_detection(file_path, {"confidence": "0.25"})
    weapon_result, _ = await run_weapon_detect(file_path, {"confidence": "0.2"})
    vehicle_result, _ = await run_vehicle_detect(file_path, {"confidence": "0.3"})
    face_result, _ = await run_face_detect(file_path, {})

    # Compile all detections
    pipeline_results = {
        "ocr": {
            "text_found": ocr_result.get("text", "")[:2000] if "error" not in ocr_result else "",
            "has_text": bool(ocr_result.get("text", "").strip()),
        },
        "objects": {
            "count": object_result.get("objects_detected", 0),
            "items": object_result.get("objects", [])[:20],
        },
        "weapons": {
            "count": weapon_result.get("weapons_detected", 0),
            "items": weapon_result.get("weapons", []),
            "threat_level": "high" if weapon_result.get("weapons_detected", 0) > 0 else "low",
        },
        "vehicles": {
            "count": vehicle_result.get("vehicles_detected", 0),
            "items": vehicle_result.get("vehicles", []),
        },
        "faces": {
            "count": face_result.get("faces_detected", 0),
            "details": face_result.get("faces", []),
        },
    }

    # Get image dimensions
    try:
        from PIL import Image
        img = Image.open(file_path)
        dimensions = {"width": img.width, "height": img.height}
        img.close()
    except Exception:
        dimensions = {"width": 0, "height": 0}

    # Generate comprehensive AI summary
    ai_report = None
    try:
        from app.ai.llm_provider import has_any_llm_key, generate_vision_base64

        if has_any_llm_key():
            import base64
            with open(file_path, "rb") as f:
                image_data = base64.b64encode(f.read()).decode("utf-8")

            ext = os.path.splitext(file_path)[1].lower()
            mime_map = {
                ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
                ".png": "image/png", ".gif": "image/gif",
                ".webp": "image/webp", ".bmp": "image/bmp",
            }
            mime_type = mime_map.get(ext, "image/jpeg")

            detection_summary = json.dumps(pipeline_results, indent=2, default=str)[:3000]

            prompt = (
                "You are an expert forensic crime scene analyst working for a police investigation unit. "
                "Analyze this crime scene image along with the automated detection results below.\n\n"
                f"Automated Detection Results:\n{detection_summary}\n\n"
                "Based on the image AND the detection data, provide a structured Crime Scene Report:\n\n"
                "## Scene Description\nDescribe what is visible in detail.\n\n"
                "## Evidence Identified\nList all potential evidence items with their locations.\n\n"
                "## Persons Detected\nDescribe any persons visible.\n\n"
                "## Vehicles\nDescribe any vehicles and their details.\n\n"
                "## Weapons/Threats\nNote any weapons or dangerous items.\n\n"
                "## Environmental Conditions\nLighting, weather, time of day if determinable.\n\n"
                "## Points of Entry/Exit\nPossible access routes.\n\n"
                "## Preliminary Assessment\nType of incident likely involved.\n\n"
                "## Recommendations\nImmediate steps for investigating officers.\n\n"
                "## Safety Concerns\nHazards for responding personnel.\n\n"
                "Be professional, factual, and thorough."
            )

            ai_report = await asyncio.to_thread(
                generate_vision_base64, image_data, mime_type, prompt, 0.3, 2048
            )

    except Exception as e:
        ai_report = f"AI analysis unavailable: {str(e)}"

    confidence = 0.75
    if pipeline_results["weapons"]["count"] > 0:
        confidence = 0.9
    elif pipeline_results["faces"]["count"] > 0:
        confidence = 0.85

    return {
        "pipeline_results": pipeline_results,
        "ai_report": ai_report,
        "image_dimensions": dimensions,
        "model_used": "YOLOv8n + InsightFace + pytesseract + AI Vision",
        "analyzed_at": datetime.utcnow().isoformat(),
        "threat_assessment": pipeline_results["weapons"]["threat_level"],
        "persons_count": pipeline_results["faces"]["count"],
        "vehicles_count": pipeline_results["vehicles"]["count"],
        "evidence_items_count": pipeline_results["objects"]["count"],
    }, confidence


# ---------------------------------------------------------------------------
# AI Summary Helper
# ---------------------------------------------------------------------------

async def generate_execution_summary(output_data: dict, tool_key: str, filename: str) -> str:
    """Generate an AI summary of tool execution results."""
    try:
        from app.ai.llm_provider import has_any_llm_key, generate_text

        if not has_any_llm_key():
            return "AI summary unavailable: No LLM API key configured."

        output_str = json.dumps(output_data, indent=2, default=str)[:5000]

        prompt = (
            "You are a forensic analyst summarizing tool execution results for a police investigation system. "
            f"Tool used: {tool_key}\n"
            f"File analyzed: {filename}\n"
            f"Results:\n{output_str}\n\n"
            "Provide a concise professional summary (3-5 sentences) of the key findings. "
            "Focus on what is most relevant for an investigation."
        )

        return await asyncio.to_thread(generate_text, prompt, 0.3, 1024)
    except Exception as e:
        return f"Summary generation unavailable: {str(e)}"


# ---------------------------------------------------------------------------
# Tool Registry
# ---------------------------------------------------------------------------

TOOL_HANDLERS: dict[str, Any] = {
    "image_ocr": run_image_ocr,
    "image_object_detect": run_object_detection,
    "image_exif": run_image_exif,
    "audio_transcribe": run_audio_transcribe,
    "document_ocr": run_document_ocr,
    "document_pdf_parse": run_document_pdf_parse,
    "digital_hash": run_digital_hash,
    "digital_metadata": run_digital_metadata,
    "digital_file_identify": run_digital_file_identify,
    "document_summarize": run_document_summarize,
    "face_detect": run_face_detect,
    "face_recognize": run_face_recognize,
    "fingerprint_match": run_fingerprint_match,
    "dna_search": run_dna_search,
    "vehicle_detect": run_vehicle_detect,
    "license_plate_ocr": run_license_plate_ocr,
    "weapon_detect": run_weapon_detect,
    "image_similarity": run_image_similarity,
    "crime_scene_analysis": run_crime_scene_analysis,
}


# ---------------------------------------------------------------------------
# Utility
# ---------------------------------------------------------------------------

def _format_size(size_bytes: int) -> str:
    """Format bytes into human-readable size."""
    if size_bytes < 1024:
        return f"{size_bytes} B"
    elif size_bytes < 1024 * 1024:
        return f"{size_bytes / 1024:.1f} KB"
    elif size_bytes < 1024 * 1024 * 1024:
        return f"{size_bytes / (1024 * 1024):.1f} MB"
    else:
        return f"{size_bytes / (1024 * 1024 * 1024):.2f} GB"
