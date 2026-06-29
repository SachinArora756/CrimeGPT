import os
from datetime import datetime
from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import settings
from app.models.case import Case
from app.models.document import DocType
from app.ai.agents.supervisor import AgentState


async def generate_legal_document(db: AsyncSession, case_id: int, doc_type: DocType, additional_context: str | None = None) -> str:
    result = await db.execute(select(Case).where(Case.id == case_id))
    case = result.scalar_one_or_none()
    if not case:
        raise ValueError(f"Case {case_id} not found")

    generators = {
        DocType.FIR: generate_fir,
        DocType.CHARGESHEET: generate_chargesheet,
        DocType.SEIZURE_MEMO: generate_seizure_memo,
        DocType.MEDICAL_LETTER: generate_medical_letter,
        DocType.COURT_LETTER: generate_court_letter,
        DocType.ARREST_MEMO: generate_arrest_memo,
    }

    generator = generators.get(doc_type)
    if not generator:
        raise ValueError(f"Unknown document type: {doc_type}")

    doc_dir = os.path.join(settings.upload_dir, str(case_id), "documents")
    os.makedirs(doc_dir, exist_ok=True)

    timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    filename = f"{doc_type.value}_{timestamp}.docx"
    file_path = os.path.join(doc_dir, filename)

    generator(case, file_path, additional_context)
    return file_path


def generate_fir(case: Case, file_path: str, context: str | None = None):
    doc = DocxDocument()

    heading = doc.add_heading("FIRST INFORMATION REPORT", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"(Under Section 154 Cr.P.C / Section 173 BNSS)")
    doc.add_paragraph("")

    table = doc.add_table(rows=12, cols=2)
    table.style = "Table Grid"

    fields = [
        ("FIR Number", case.fir_number),
        ("Date & Time of FIR", datetime.utcnow().strftime("%d/%m/%Y %H:%M")),
        ("Police Station", case.station_id or "—"),
        ("Complainant Name", case.complainant_name),
        ("Complainant Address", case.complainant_address or "—"),
        ("Complainant Contact", case.complainant_contact or "—"),
        ("Date of Occurrence", str(case.incident_date) if case.incident_date else "—"),
        ("Time of Occurrence", case.incident_time or "—"),
        ("Place of Occurrence", case.incident_location or "—"),
        ("Name of Accused", case.accused_name or "Unknown"),
        ("Offense Type", case.offense_type or "—"),
        ("Sections Applied", ", ".join(case.sections_applied) if case.sections_applied else "—"),
    ]

    for i, (label, value) in enumerate(fields):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = str(value)
        for cell in row.cells:
            for para in cell.paragraphs:
                para.style.font.size = Pt(10)

    doc.add_paragraph("")
    doc.add_heading("Details of Complaint:", level=2)
    doc.add_paragraph(case.description)

    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("_" * 40)
    doc.add_paragraph("Signature of Complainant")
    doc.add_paragraph("")
    doc.add_paragraph("_" * 40)
    doc.add_paragraph("Signature of Officer Recording FIR")

    doc.save(file_path)


def generate_chargesheet(case: Case, file_path: str, context: str | None = None):
    doc = DocxDocument()

    heading = doc.add_heading("CHARGE SHEET", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("(Under Section 173 CrPC / Section 193 BNSS)")
    doc.add_paragraph("")

    doc.add_paragraph(f"FIR Number: {case.fir_number}")
    doc.add_paragraph(f"Police Station: {case.station_id or '—'}")
    doc.add_paragraph(f"Date of Filing: {datetime.utcnow().strftime('%d/%m/%Y')}")
    doc.add_paragraph("")

    doc.add_heading("1. Particulars of the Accused:", level=2)
    doc.add_paragraph(f"Name: {case.accused_name or 'Unknown'}")
    doc.add_paragraph("")

    doc.add_heading("2. Offense:", level=2)
    doc.add_paragraph(f"Type: {case.offense_type or '—'}")
    doc.add_paragraph(f"Sections: {', '.join(case.sections_applied) if case.sections_applied else '—'}")
    doc.add_paragraph("")

    doc.add_heading("3. Brief Facts:", level=2)
    doc.add_paragraph(case.description)
    doc.add_paragraph("")

    doc.add_heading("4. Evidence:", level=2)
    doc.add_paragraph("(As per attached evidence list)")
    doc.add_paragraph("")

    doc.add_heading("5. List of Witnesses:", level=2)
    doc.add_paragraph("1. Complainant: " + case.complainant_name)
    doc.add_paragraph("")

    doc.add_paragraph("")
    doc.add_paragraph("_" * 40)
    doc.add_paragraph("Investigating Officer")

    doc.save(file_path)


def generate_seizure_memo(case: Case, file_path: str, context: str | None = None):
    doc = DocxDocument()

    heading = doc.add_heading("SEIZURE MEMO", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Case FIR: {case.fir_number}")
    doc.add_paragraph(f"Date: {datetime.utcnow().strftime('%d/%m/%Y')}")
    doc.add_paragraph(f"Place of Seizure: {case.incident_location or '—'}")
    doc.add_paragraph("")

    doc.add_heading("Articles Seized:", level=2)
    doc.add_paragraph("(To be filled by the investigating officer)")
    doc.add_paragraph("")

    doc.add_heading("Witnesses:", level=2)
    doc.add_paragraph("1. ________________________")
    doc.add_paragraph("2. ________________________")
    doc.add_paragraph("")

    doc.add_paragraph("_" * 40)
    doc.add_paragraph("Investigating Officer")

    doc.save(file_path)


def generate_medical_letter(case: Case, file_path: str, context: str | None = None):
    doc = DocxDocument()

    heading = doc.add_heading("LETTER FOR MEDICAL EXAMINATION", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Date: {datetime.utcnow().strftime('%d/%m/%Y')}")
    doc.add_paragraph(f"FIR No: {case.fir_number}")
    doc.add_paragraph("")

    doc.add_paragraph("To,")
    doc.add_paragraph("The Medical Officer,")
    doc.add_paragraph("[Hospital Name]")
    doc.add_paragraph("")

    doc.add_paragraph("Subject: Request for Medical Examination")
    doc.add_paragraph("")
    doc.add_paragraph(
        f"Sir/Madam, kindly conduct medical examination of "
        f"{case.complainant_name} in connection with FIR No. {case.fir_number} "
        f"registered at this police station."
    )
    doc.add_paragraph("")
    doc.add_paragraph("_" * 40)
    doc.add_paragraph("Station House Officer")

    doc.save(file_path)


def generate_court_letter(case: Case, file_path: str, context: str | None = None):
    doc = DocxDocument()

    heading = doc.add_heading("LETTER TO COURT", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Date: {datetime.utcnow().strftime('%d/%m/%Y')}")
    doc.add_paragraph(f"FIR No: {case.fir_number}")
    doc.add_paragraph("")

    doc.add_paragraph("To,")
    doc.add_paragraph("The Hon'ble Court of [Magistrate],")
    doc.add_paragraph("[Court Address]")
    doc.add_paragraph("")

    doc.add_paragraph("Subject: Submission of Charge Sheet / Application for Remand")
    doc.add_paragraph("")
    doc.add_paragraph(
        f"Respectfully submitted that in connection with FIR No. {case.fir_number}, "
        f"investigation has been conducted and the following is submitted for your kind consideration."
    )
    doc.add_paragraph("")
    doc.add_paragraph(f"Brief facts: {case.description[:300]}")
    doc.add_paragraph("")
    doc.add_paragraph("_" * 40)
    doc.add_paragraph("Investigating Officer")

    doc.save(file_path)


def generate_arrest_memo(case: Case, file_path: str, context: str | None = None):
    doc = DocxDocument()

    heading = doc.add_heading("ARREST MEMO", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("(Under Section 41B CrPC / Section 35 BNSS)")
    doc.add_paragraph("")

    doc.add_paragraph(f"FIR No: {case.fir_number}")
    doc.add_paragraph(f"Date of Arrest: {datetime.utcnow().strftime('%d/%m/%Y')}")
    doc.add_paragraph(f"Time of Arrest: {datetime.utcnow().strftime('%H:%M')}")
    doc.add_paragraph("")

    doc.add_heading("Particulars of Arrested Person:", level=2)
    doc.add_paragraph(f"Name: {case.accused_name or '—'}")
    doc.add_paragraph("")

    doc.add_heading("Grounds of Arrest:", level=2)
    doc.add_paragraph(f"Offense: {case.offense_type or '—'}")
    doc.add_paragraph(f"Sections: {', '.join(case.sections_applied) if case.sections_applied else '—'}")
    doc.add_paragraph("")

    doc.add_paragraph("Rights of the Arrested Person were communicated: Yes / No")
    doc.add_paragraph("")
    doc.add_paragraph("_" * 40)
    doc.add_paragraph("Arresting Officer")
    doc.add_paragraph("")
    doc.add_paragraph("_" * 40)
    doc.add_paragraph("Witness 1")

    doc.save(file_path)


def document_gen_node(state: AgentState) -> AgentState:
    state["current_agent"] = "document_gen"
    return state
